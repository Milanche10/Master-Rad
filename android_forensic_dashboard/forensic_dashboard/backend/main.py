"""
main.py — FastAPI backend za Android Forensic Dashboard
────────────────────────────────────────────────────────
Endpointi:
  POST /api/session                       → otvori dump, vrati session_id
  GET  /api/session/{id}                  → info o sesiji
  DELETE /api/session/{id}               → obriši sesiju
  GET  /api/session/{id}/analyze/{module} → pokreni jedan modul
  POST /api/session/{id}/analyze/all      → pokreni sve module
  GET  /api/session/{id}/correlations     → cross-reference analiza
  GET  /api/session/{id}/timeline         → hronološka linija artefakata
  GET  /api/session/{id}/report           → generiši izveštaj
"""

import io
import sys
import subprocess
import uuid
import logging
import asyncio
import hashlib
import traceback
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from datetime import datetime, timezone
from xml.sax.saxutils import escape as _esc

TOOL_VERSION = "1.0.0"

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, FileResponse
from pydantic import BaseModel

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable,
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Registruj UTF-8 font koji podržava srpska slova (š đ č ć ž).
# Probamo redom: Arial (Windows), DejaVuSans (Linux/pip), Liberation (Linux).
_PDF_FONT = "Helvetica"   # fallback ako nijedan TTF ne nađemo
_PDF_FONT_BOLD = "Helvetica-Bold"
for _candidate, _bold, _alias, _alias_bold in [
    (r"C:\Windows\Fonts\arial.ttf",   r"C:\Windows\Fonts\arialbd.ttf",   "ArialUTF", "ArialUTF-Bold"),
    ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
     "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", "DejaVuSans", "DejaVuSans-Bold"),
    ("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
     "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf", "Liberation", "Liberation-Bold"),
]:
    if Path(_candidate).exists():
        try:
            pdfmetrics.registerFont(TTFont(_alias, _candidate))
            if Path(_bold).exists():
                pdfmetrics.registerFont(TTFont(_alias_bold, _bold))
                _PDF_FONT_BOLD = _alias_bold
            _PDF_FONT = _alias
            break
        except Exception:
            pass
from docx import Document

# Moduli analize
from modules import (
    device_info, sms, calllog, contacts, browser,
    wifi, apk, exif, crypto, mp3_signal, blockchain, signal_brd,
    anti_forensics, notes, reminders, app_messaging, deleted_recovery,
)
from utils.dump_resolver import DumpResolver
from utils.entity_graph import build_entity_graph
from utils.db_inventory import scan_all_databases, inventory_summary
from utils import evidence, case_store, audit_log
from report_engine import ForensicReportEngine
from report.report_model import build_report_model
from correlation.rules_registry import registry as correlation_registry
from ai_analyst import generate_ai_conclusion, ai_available

# ─── Acquisition & Export slojevi (novi; ne diraju analitički engine) ──────
from acquisition import detect as acq_detect
from acquisition import jobs as acq_jobs
from acquisition import cases_fs as acq_cases
from acquisition import storage as acq_storage
from export import exporters as exporters_mod
from export import packager as packager_mod

# ─── Logging ──────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)-7s  %(name)s  %(message)s",
)
logger = logging.getLogger("afd")

# ─── App ──────────────────────────────────────────────────────────────────
app = FastAPI(
    title="Android Forensic Dashboard API",
    version="1.0.0",
    docs_url="/api/docs",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://127.0.0.1:3000"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ─── In-memory session store ───────────────────────────────────────────────
# { session_id: { "dump_path": str, "resolver": DumpResolver, "results": {}, "created_at": str } }
SESSIONS: dict = {}

# ─── Mapiranje naziva modula → funkcija ───────────────────────────────────
MODULE_MAP = {
    "device_info":  device_info.analyze,
    "sms":          sms.analyze,
    "calllog":      calllog.analyze,
    "contacts":     contacts.analyze,
    "browser":      browser.analyze,
    "wifi":         wifi.analyze,
    "apk":          apk.analyze,
    "exif":         exif.analyze,
    "crypto":       crypto.analyze,
    "mp3_signal":   mp3_signal.analyze,
    "blockchain":   blockchain.analyze,
    "signal_brd":   signal_brd.analyze,
    "app_messaging": app_messaging.analyze,
    "notes":        notes.analyze,
    "reminders":    reminders.analyze,
    "deleted_recovery": deleted_recovery.analyze,
    "anti_forensics": anti_forensics.analyze,
}

MODULE_ORDER = list(MODULE_MAP.keys())


# ─── Pydantic modeli ──────────────────────────────────────────────────────
class OpenDumpRequest(BaseModel):
    dump_path: str
    examiner: str = ""       # ime veštaka (chain of custody); opciono
    fs_case_id: str = ""     # ako dolazi iz akvizicije (Case_YYYY_NNNN) — poveži slučaj
    source: str = "dump"     # izvor dokaza: dump | mobile | sim | sdcard | usb


class RevealRequest(BaseModel):
    path: str


class ArtifactExportRequest(BaseModel):
    artifact: dict


class AcquireRequest(BaseModel):
    examiner: str = ""
    # SD/USB
    mount: str = ""
    disk_info: dict = {}
    # telefon
    serial: str = ""
    device_info: dict = {}
    # SIM
    reader: str = ""


# ─── Helpers ──────────────────────────────────────────────────────────────
def get_session(session_id: str) -> dict:
    if session_id not in SESSIONS:
        raise HTTPException(status_code=404, detail="Sesija nije pronađena")
    return SESSIONS[session_id]


def resolve_artifact_file(session: dict, rel_path: str) -> Path:
    """
    Pretvara 'source' polje artefakta (relativna putanja unutar dump-a) u
    apsolutnu putanju, sa zaštitom od path traversal-a. Baca HTTPException
    ako fajl ne postoji ili je van dump root-a.
    """
    resolver: DumpResolver = session["resolver"]
    root = resolver.root.resolve()

    rel_path = (rel_path or "").strip().lstrip("/\\")
    candidate = (root / rel_path).resolve()

    if root not in candidate.parents and candidate != root:
        raise HTTPException(status_code=400, detail="Putanja je van dump-a")

    if not candidate.exists() or not candidate.is_file():
        raise HTTPException(status_code=404, detail=f"Fajl nije pronađen: {rel_path}")

    return candidate


def _normalize_ts(ts: str) -> str:
    """Normalizuje timestamp u jedinstveni ISO 8601 UTC format (…Z)."""
    dt = _parse_ts(ts)
    if not dt:
        return ts
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=timezone.utc)
    return dt.astimezone(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _event_severity(e: dict) -> str:
    """
    Nivo ozbiljnosti događaja za timeline/izveštaj:
    high   — šifrovana komunikacija, trojanizovane aplikacije, stego, covert signaling
    medium — pozivi, lokacije, kripto, nalozi, kontakti, komunikacija
    low    — web istorija, generička media i ostalo
    """
    extra = e.get("extra") or {}
    if extra.get("encryption") or extra.get("trojanized") or extra.get("suspicious"):
        return "high"
    if extra.get("call_type") == "zero_second_signal" or extra.get("sideloaded"):
        return "high"
    if e.get("type") in ("call", "location", "crypto", "account", "comm", "contact"):
        return "medium"
    return "low"


def _web_host(value: str) -> str:
    """Izvuci host iz web događaja tipa 'www.nzz.ch – Naslov strane'."""
    if not value:
        return ""
    head = value.split(" – ")[0].split(" - ")[0].strip()
    # uzmi samo domen (prvi token bez razmaka)
    return head.split()[0] if head else ""


WEB_BURST_MIN = 3   # koliko uzastopnih web događaja se sažima u jedan red


def collapse_timeline_for_report(timeline: list) -> list:
    """
    Pripremi čitljive redove za izveštaj: uzastopni 'web' događaji (browser
    istorija — reddit, imdb, apkpure...) sažimaju se u JEDAN sažeti red sa
    vremenskim opsegom i top hostovima. Ostali događaji ostaju pojedinačni.
    Time se detaljna vremenska linija svodi sa stotina redova šuma na
    preglednu, forenzički relevantnu hronologiju.
    """
    rows = []
    i, n = 0, len(timeline)
    while i < n:
        e = timeline[i]
        if e.get("type") == "web":
            j = i
            hosts = []
            while j < n and timeline[j].get("type") == "web":
                h = _web_host(timeline[j].get("value", ""))
                if h and h not in hosts:
                    hosts.append(h)
                j += 1
            run = timeline[i:j]
            if len(run) >= WEB_BURST_MIN:
                top = ", ".join(hosts[:4]) + (f" … (+{len(hosts) - 4})" if len(hosts) > 4 else "")
                rows.append({
                    "kind": "burst",
                    "ts": run[0].get("ts"),
                    "ts_end": run[-1].get("ts"),
                    "type": "web",
                    "value": f"{len(run)} web poseta: {top}",
                    "source": "Chrome/History",
                    "severity": "low",
                })
                i = j
                continue
        rows.append({"kind": "event", **e})
        i += 1
    return rows


def build_timeline(results: dict) -> list:
    events = []
    seen_keys = set()
    for module_name, data in results.items():
        for artifact in (data.get("artifacts") or []):
            if artifact.get("ts"):
                key = (module_name, artifact.get("type"), artifact.get("value"), artifact.get("ts"))
                if key in seen_keys:
                    continue
                seen_keys.add(key)
                events.append({
                    **artifact,
                    "ts": _normalize_ts(artifact["ts"]),
                    "module": module_name,
                    "severity": _event_severity(artifact),
                })
    return sorted(events, key=lambda e: e["ts"])


# Tipovi/uslovi artefakata koji predstavljaju "događaje" relevantne za
# rekonstrukciju vremenske linije slučaja (headline timeline). Ideja je da
# headline timeline priča priču ("poziv → lokacija → poruka → ...") bez
# šuma generičke browser istorije, svake pojedinačne SMS poruke itd.
def _is_headline_event(e: dict) -> bool:
    etype = e.get("type")
    extra = e.get("extra") or {}

    if etype in ("call", "location"):
        return True
    if etype == "crypto":
        return True
    if etype == "comm":
        # samo šifrovane/sumnjive poruke, ne sve SMS
        return bool(extra.get("encryption"))
    if etype == "app":
        return bool(extra.get("trojanized") or extra.get("sideloaded"))
    if etype == "media":
        return bool(extra.get("suspicious"))
    if etype == "account":
        return True
    if etype == "web":
        # samo visokorizične kategorije (crypto exchange, VPN, dark web...)
        return bool(extra.get("categories"))
    return False


def build_headline_timeline(results: dict, correlations: list | None = None) -> list:
    """
    Sažeta vremenska linija "glavnih događaja" — za rekonstrukciju
    redosleda radnji (poziv → lokacija → transakcija → ...), bez šuma
    pojedinačnih SMS poruka i browser istorije.
    """
    detailed = build_timeline(results)
    headline = [e for e in detailed if _is_headline_event(e)]

    # Dodaj i artefakte direktno povezane sa korelacijama (čak i ako po
    # default kriterijumu ne bi ušli) — korisno za "veza A → B" priču.
    if correlations:
        seen = {(e.get("module"), e.get("value"), e.get("ts")) for e in headline}
        for c in correlations:
            for a in c.get("linked_artifacts") or []:
                if not a.get("ts"):
                    continue
                key = (a.get("module"), a.get("value"), a.get("ts"))
                if key not in seen:
                    seen.add(key)
                    headline.append(a)

    return sorted(headline, key=lambda e: e["ts"])


def _parse_ts(ts):
    if not ts:
        return None
    try:
        return datetime.fromisoformat(ts.replace("Z", "+00:00"))
    except Exception:
        return None


def detect_timeline_anomalies(timeline: list) -> dict:
    """
    Temporalne anomalije (#3): buduće/pre-Android/epoch datumi, neparsabilni
    timestamp-ovi, i praznine (gaps) veće od 72h u aktivnosti. Sve iz postojećih
    ts vrednosti — bez izmišljanja.
    """
    now = datetime.now(timezone.utc)
    android_epoch = datetime(2008, 9, 1, tzinfo=timezone.utc)  # pre Androida = nemoguće
    future, pre_android, unparsed = [], [], 0
    parsed = []
    for e in timeline:
        dt = _parse_ts(e.get("ts"))
        if dt is None:
            unparsed += 1
            continue
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        if dt > now:
            future.append(e.get("ts"))
        elif dt < android_epoch:
            pre_android.append(e.get("ts"))
        parsed.append(dt)

    # praznine > 72h između uzastopnih događaja
    parsed.sort()
    gaps = []
    for a, b in zip(parsed, parsed[1:]):
        h = (b - a).total_seconds() / 3600
        if h > 72:
            gaps.append({"od": a.strftime("%Y-%m-%dT%H:%M:%SZ"),
                         "do": b.strftime("%Y-%m-%dT%H:%M:%SZ"), "sati": round(h, 1)})

    return {
        "future_dated": len(future),
        "pre_android": len(pre_android),
        "unparsed": unparsed,
        "large_gaps": gaps[:10],
        "gap_count": len(gaps),
        "anomaly_total": len(future) + len(pre_android) + unparsed,
    }


def _enrich_module(session: dict, module_name: str, result: dict) -> dict:
    """Normalization layer: obogati artefakte jednog modula (hash/provenance/id/confidence)."""
    resolver = session.get("resolver")
    case_id = session.get("case_id")
    try:
        arts = result.get("artifacts") or []
        result = dict(result)
        result["artifacts"] = [evidence.enrich_artifact(a, module_name, resolver, case_id) for a in arts]
    except Exception as e:
        logger.warning(f"Enrichment failed for {module_name}: {e}")
    return result


def _results_content_hash(results: dict) -> str:
    """
    Deterministički SHA-256 nad SADRŽAJEM nalaza (bez vremena generisanja/uuid-a).
    Isti dump → isti heš → reproducibilnost. Ne uključuje volatilne vrednosti.
    """
    canon = []
    for module in sorted(results.keys()):
        data = results[module]
        arts = sorted(
            [(a.get("type"), a.get("value"), a.get("source"), a.get("ts")) for a in (data.get("artifacts") or [])]
        )
        alerts = sorted(data.get("alerts") or [])
        canon.append((module, data.get("status"), arts, alerts))
    blob = repr(canon).encode("utf-8", "replace")
    return hashlib.sha256(blob).hexdigest()


def _dump_fingerprint(resolver) -> str:
    """Otisak ulaza (root + broj i imena SQLite baza) za reproducibilnost."""
    try:
        idx = resolver._build_sqlite_index()
        names = sorted(p.name for p in idx.keys())
        basis = f"{resolver.root}|{len(names)}|{'|'.join(names[:100])}"
        return hashlib.sha256(basis.encode("utf-8", "replace")).hexdigest()[:32]
    except Exception:
        return ""


def build_correlations(results: dict) -> list:
    """
    Cross-reference engine — pronalazi zajedničke entitete
    između artefakata iz različitih modula.
    """
    correlations = []

    def artifacts_of(module: str) -> list:
        return results.get(module, {}).get("artifacts", [])

    def find_in(artifacts, **kwargs) -> list:
        """Vrati artefakte koji u extra polju imaju date vrednosti."""
        matches = []
        for a in artifacts:
            extra = a.get("extra") or {}
            if all(extra.get(k) == v for k, v in kwargs.items()):
                matches.append(a)
        return matches

    def _mk_corr(cid, title, sources, detail, linked, score, factors=None):
        """
        Konstruiši korelaciju sa NUMERIČKIM skorom (0-100) i EKSPLICITNIM
        dokazima. Invarijanta: korelacija bez citiranog artefakta se ne pravi
        (osim ako score sam nosi obrazloženje). 'confidence' string je alias
        opsega radi kompatibilnosti sa postojećim izveštajem/frontendom.
        """
        score = max(0, min(100, int(score)))
        band = "VISOKA" if score >= 70 else ("SREDNJA" if score >= 40 else "NISKA")
        evidence = [{
            "artifact_id": a.get("id"),
            "module": a.get("module"),
            "source": a.get("source"),
            "value": (a.get("value") or "")[:80],
        } for a in (linked or [])]
        return {
            "id": cid, "title": title,
            "score": score, "band": band, "confidence": band,  # confidence = alias
            "sources": sources, "detail": detail,
            "scoring_factors": factors or [],
            "evidence": evidence,
            "linked_artifacts": linked or [],
        }

    # ── C1: Isti broj telefona u SMS + calllog + contacts ─────────────────
    sms_artifacts   = artifacts_of("sms")
    call_artifacts  = artifacts_of("calllog")
    cont_artifacts  = artifacts_of("contacts")

    # DEFENZIVNI pristup (.get umesto subscript) — artefakt bez 'phone' ne sme da obori
    sms_phones   = set(a.get("extra", {}).get("phone") for a in sms_artifacts if a.get("extra", {}).get("phone"))
    call_phones  = set(a.get("extra", {}).get("phone") for a in call_artifacts if a.get("extra", {}).get("phone"))
    cont_phones  = set(a.get("extra", {}).get("phone") for a in cont_artifacts if a.get("extra", {}).get("phone"))

    common_phones = sms_phones & call_phones
    for phone in sorted(common_phones)[:5]:   # sorted → deterministički izbor
        in_contacts = phone in cont_phones
        has_encrypted = any(
            a.get("extra", {}).get("encryption")
            for a in sms_artifacts
            if a.get("extra", {}).get("phone") == phone
        )
        has_zero_sec = any(
            a.get("extra", {}).get("call_type") == "zero_second_signal"
            for a in call_artifacts
            if a.get("extra", {}).get("phone") == phone
        )

        sources = ["SMS (mmssms.db)", "Calllog (calllog.db)"]
        if in_contacts:
            sources.append("Contacts (contacts2.db)")

        detail = f"Broj {phone} se pojavljuje u {len(sources)} nezavisnih izvora."
        if has_encrypted:
            detail += " SMS poruke su šifrovane (AES/HHY pattern)."
        if has_zero_sec:
            detail += " Detektovani nultosekundni pozivi ka ovom broju."

        linked = (
            [a for a in sms_artifacts if a.get("extra", {}).get("phone") == phone][:2] +
            [a for a in call_artifacts if a.get("extra", {}).get("phone") == phone][:2]
        )
        # Skoring: baza 45 (dva izvora) + kontakt +10 + jaki indikatori +20/+15
        score = 45 + (10 if in_contacts else 0) + (20 if has_encrypted else 0) + (15 if has_zero_sec else 0)
        factors = [f"{len(sources)} nezavisna izvora"]
        if in_contacts: factors.append("broj u kontaktima")
        if has_encrypted: factors.append("šifrovane poruke")
        if has_zero_sec: factors.append("nultosekundni pozivi")
        correlations.append(_mk_corr(
            f"C-TEL-{phone[-4:]}", f"Komunikaciona korelacija: {phone}",
            sources, detail, linked, score, factors))

    # ── C2: OmniNotes creator + šifrovane poruke ─────────────────────────
    # ISPRAVKA (forenzička validnost): NE tvrdimo detalje DEX analize (klasa/algoritam)
    # osim ako ih APK modul zaista prijavi kao artefakt. Navodimo samo ono što
    # citirani artefakti podržavaju.
    omninotes_alerts = [
        a for a in results.get("sms", {}).get("alerts", [])
        if "omninotes" in a.lower()
    ]
    encrypted_sms = [a for a in sms_artifacts if a.get("extra", {}).get("encryption")]
    apk_trojan = [a for a in artifacts_of("apk") if a.get("extra", {}).get("trojanized")]
    if omninotes_alerts and encrypted_sms:
        detail = (
            f"OmniNotes je zabeležen kao SMS creator uz {len(encrypted_sms)} šifrovanih poruka "
            f"(mmssms.db creator kolona)."
        )
        if apk_trojan:
            cats = sorted({c for a in apk_trojan for c in (a.get("extra", {}).get("categories") or [])})
            detail += f" APK statička analiza potvrđuje neočekivane DEX reference: {', '.join(cats) or 'SMS/kripto'}."
            score = 85
        else:
            detail += " (APK modul nije potvrdio trojanizaciju — potrebna statička DEX analiza za punu potvrdu.)"
            score = 60
        correlations.append(_mk_corr(
            "C-APK-SMS", "Trojanizovana aplikacija → šifrovane SMS poruke",
            ["SMS creator (mmssms.db)"] + (["APK/DEX analiza"] if apk_trojan else []),
            detail, encrypted_sms[:3] + apk_trojan[:2], score,
            ["OmniNotes kao SMS creator", f"{len(encrypted_sms)} šifrovanih poruka"] +
            (["APK trojanizacija potvrđena"] if apk_trojan else ["APK nepotvrđen"])))

    # ── C3: Browser kriptovalute ──────────────────────────────────────────
    # ISPRAVKA: uklonjena necitirana tvrdnja o akreditivima u Login Data.
    browser_crypto = [
        a for a in artifacts_of("browser")
        if "crypto" in " ".join(a.get("extra", {}).get("categories", [])).lower()
        or "P2P" in " ".join(a.get("extra", {}).get("categories", []))
    ]
    login_artifacts = [a for a in artifacts_of("browser") if a.get("extra", {}).get("credential")
                       or "Login Data" in (a.get("source") or "")]
    if browser_crypto:
        detail = f"Pronađeno {len(browser_crypto)} poseta kriptovalutnim/P2P exchange servisima (Chrome History)."
        score = 45
        if login_artifacts:
            detail += f" Dodatno {len(login_artifacts)} sačuvanih akreditiva (Login Data)."
            score = 55
        correlations.append(_mk_corr(
            "C-CRYPTO-WEB", "Kriptovalutne aktivnosti u browser istoriji",
            ["Chrome History"] + (["Chrome Login Data"] if login_artifacts else []),
            detail, browser_crypto[:3] + login_artifacts[:2], score,
            [f"{len(browser_crypto)} kripto/P2P poseta"]))

    # ── C4: Međunarodna komunikacija ──────────────────────────────────────
    # ISPRAVKA: NE emitujemo korelaciju bez citiranog artefakta. Gradimo dokaz
    # iz stvarnih SMS/calllog artefakata sa stranim brojem, ne iz alert stringa.
    intl_sms = [a for a in sms_artifacts if a.get("extra", {}).get("country")
                and a.get("extra", {}).get("country") != "Srbija"]
    intl_call = [a for a in call_artifacts if a.get("extra", {}).get("country")
                 and a.get("extra", {}).get("country") != "Srbija"]
    intl_evidence = intl_sms + intl_call
    if intl_evidence:
        countries = sorted({a.get("extra", {}).get("country") for a in intl_evidence if a.get("extra", {}).get("country")})
        detail = (f"Međunarodna komunikacija: {len(intl_sms)} SMS + {len(intl_call)} poziva "
                  f"ka {len(countries)} država ({', '.join(countries)}).")
        correlations.append(_mk_corr(
            "C-INTL", "Međunarodna komunikaciona mreža",
            ["SMS (mmssms.db)", "Calllog (calllog.db)"], detail,
            intl_evidence[:4], 40 + min(20, len(countries) * 10),
            [f"{len(countries)} strana država", f"{len(intl_evidence)} događaja"]))

    # ── C5: Lokacija — EXIF foto vremenski blizu WiFi konekcije ───────────
    GEO_TIME_WINDOW_HOURS = 6
    exif_locations = [a for a in artifacts_of("exif") if a.get("ts") and a.get("type") == "location"]
    wifi_networks = [a for a in artifacts_of("wifi") if a.get("ts")]

    for photo in exif_locations:
        photo_ts = _parse_ts(photo.get("ts"))
        if not photo_ts:
            continue
        for net in wifi_networks:
            net_ts = _parse_ts(net.get("ts"))
            if not net_ts:
                continue
            delta_h = abs((photo_ts - net_ts).total_seconds()) / 3600
            if delta_h <= GEO_TIME_WINDOW_HOURS:
                ssid = net.get("extra", {}).get("ssid", "?")
                correlations.append(_mk_corr(
                    f"C-GEO-{ssid[:10]}", f"Fizička lokacija: foto + WiFi mreža \"{ssid}\"",
                    ["EXIF/GPS (DCIM)", "WiFi (WifiConfigStore.xml)"],
                    (f"Fotografija sa GPS koordinatama snimljena je {delta_h:.1f}h od trenutka "
                     f"poslednje konekcije na mrežu \"{ssid}\" — ukazuje na fizičku lokaciju uređaja."),
                    [photo, net], 85 if delta_h <= 1 else 60,
                    [f"vremenska razlika {delta_h:.1f}h", "GPS + WiFi"]))
                break  # jedan par po fotografiji je dovoljan

    # ── C6: APK statička analiza ↔ šifrovane SMS poruke ────────────────────
    apk_trojan = [a for a in artifacts_of("apk") if a.get("extra", {}).get("trojanized")]
    if apk_trojan and encrypted_sms:
        correlations.append(_mk_corr(
            "C-APK-DEX-SMS",
            "Statička DEX analiza potvrđuje trojanizovanu aplikaciju kao izvor šifrovanih SMS poruka",
            ["APK/DEX statička analiza", "SMS (mmssms.db)"],
            (f"APK modul je identifikovao {len(apk_trojan)} modifikovanu aplikaciju sa "
             f"neočekivanim DEX referencama, što se poklapa sa {len(encrypted_sms)} šifrovanih SMS poruka."),
            apk_trojan[:2] + encrypted_sms[:2], 85,
            ["DEX trojanizacija", f"{len(encrypted_sms)} šifrovanih poruka"]))

    # ── C7: Finansijski trag — ista kripto adresa kroz više modula ────────
    addr_sources = {}
    for module_name in ("crypto", "blockchain", "signal_brd", "browser"):
        for a in artifacts_of(module_name):
            addr = a.get("extra", {}).get("address") or a.get("extra", {}).get("address_or_uri")
            if addr:
                addr_sources.setdefault(addr, []).append((module_name, a))

    for addr, occurrences in addr_sources.items():
        modules_involved = sorted(set(m for m, _ in occurrences))
        if len(modules_involved) >= 2:
            correlations.append(_mk_corr(
                f"C-CRYPTO-FLOW-{addr[-6:]}",
                f"Kriptovalutni finansijski trag: adresa {addr[:10]}...{addr[-6:]}",
                [f"Modul: {m}" for m in modules_involved],
                (f"Adresa {addr} se pojavljuje u {len(modules_involved)} nezavisnih modula "
                 f"({', '.join(modules_involved)}) — rekonstrukcija finansijskog toka."),
                [a for _, a in occurrences[:4]], 85 if len(modules_involved) >= 3 else 60,
                [f"{len(modules_involved)} modula", "ista kripto adresa"]))

    # ── C8: Covert signaling — audio steganografija + nultosekundni pozivi ─
    AUDIO_TIME_WINDOW_HOURS = 2
    suspicious_audio = [a for a in artifacts_of("mp3_signal") if a.get("extra", {}).get("suspicious")]
    zero_sec_calls = [a for a in call_artifacts if a.get("extra", {}).get("call_type") == "zero_second_signal"]

    if suspicious_audio and zero_sec_calls:
        correlations.append(_mk_corr(
            "C-STEGO-SIGNAL", "Audio steganografija u vremenskoj blizini nultosekundnih poziva",
            ["MP3/Audio analiza", "Calllog (calllog.db)"],
            (f"Pronađeno {len(suspicious_audio)} audio fajl(ova) sa indikatorima steganografije i "
             f"{len(zero_sec_calls)} nultosekundnih poziva — mogući prikriveni signalni kanal."),
            suspicious_audio[:2] + zero_sec_calls[:2], 55,
            [f"{len(suspicious_audio)} sumnjivih audio", f"{len(zero_sec_calls)} nultosek. poziva"]))

    # ── C9: Generička korelacija — zajednički identifikator kroz module ───
    # Radi za bilo koji dump: traži email/username/account/imei/device_id/
    # serial koji se pojavljuje u extra poljima artefakata iz 2+ modula.
    IDENTITY_KEYS = ("email", "username", "account", "account_id", "imei", "device_id", "serial", "uid")
    identity_sources = {}
    for module_name, module_result in results.items():
        for a in module_result.get("artifacts", []):
            extra = a.get("extra") or {}
            for key in IDENTITY_KEYS:
                val = extra.get(key)
                if val and isinstance(val, str) and len(val) >= 4:
                    identity_sources.setdefault((key, val), set()).add(module_name)

    seen_ids = 0
    for (id_key, id_val), modules_involved in identity_sources.items():
        if len(modules_involved) < 2:
            continue
        if seen_ids >= 8:
            break
        seen_ids += 1
        linked = []
        for module_name in modules_involved:
            for a in results.get(module_name, {}).get("artifacts", []):
                extra = a.get("extra") or {}
                if extra.get(id_key) == id_val:
                    linked.append(a)
                    break
        correlations.append(_mk_corr(
            f"C-ID-{id_key.upper()}-{id_val[-6:]}",
            f"Zajednički identifikator ({id_key}): {id_val}",
            [f"Modul: {m}" for m in sorted(modules_involved)],
            (f"Vrednost \"{id_val}\" (polje: {id_key}) pojavljuje se u "
             f"{len(modules_involved)} nezavisnih modula ({', '.join(sorted(modules_involved))})."),
            linked[:4], 80 if len(modules_involved) >= 3 else 55,
            [f"{len(modules_involved)} modula", f"isti {id_key}"]))

    # ── C10: Generička lokaciona korelacija — bilo koja dva "location"
    # artefakta iz različitih modula u vremenskoj blizini ─────────────────
    LOCATION_TIME_WINDOW_HOURS = 6
    all_locations = []
    for module_name, module_result in results.items():
        for a in module_result.get("artifacts", []):
            if a.get("type") == "location" and a.get("ts"):
                all_locations.append((module_name, a))

    used_pairs = set()
    geo_count = 0
    for i, (mod_a, a) in enumerate(all_locations):
        ts_a = _parse_ts(a.get("ts"))
        if not ts_a:
            continue
        for mod_b, b in all_locations[i + 1:]:
            if mod_b == mod_a:
                continue
            ts_b = _parse_ts(b.get("ts"))
            if not ts_b:
                continue
            delta_h = abs((ts_a - ts_b).total_seconds()) / 3600
            if delta_h <= LOCATION_TIME_WINDOW_HOURS:
                pair_key = (mod_a, a.get("value"), mod_b, b.get("value"))
                if pair_key in used_pairs:
                    continue
                used_pairs.add(pair_key)
                geo_count += 1
                if geo_count > 8:
                    break
                correlations.append(_mk_corr(
                    f"C-GEOGEN-{mod_a}-{mod_b}-{geo_count}",
                    f"Vremenska i lokacijska korelacija: {mod_a} ↔ {mod_b}",
                    [f"Modul: {mod_a}", f"Modul: {mod_b}"],
                    (f"Lokacijski događaji iz modula \"{mod_a}\" i \"{mod_b}\" zabeleženi su "
                     f"{delta_h:.1f}h jedan od drugog — fizička lokacija uređaja."),
                    [a, b], 85 if delta_h <= 1 else 60,
                    [f"vremenska razlika {delta_h:.1f}h", "dva lokacijska izvora"]))
        if geo_count > 8:
            break

    # Dedup: ukloni korelacije sa ISTIM skupom dokaza i naslovom (zadrži onu
    # sa najvišim skorom). Sprečava dupliranje generičkih pravila (C9/C10) i
    # višestruko izveštavanje istog nalaza.
    seen_ev, deduped = set(), []
    for c in sorted(correlations, key=lambda c: -c.get("score", 0)):
        ev_ids = tuple(sorted(
            (e.get("artifact_id") or e.get("value") or "") for e in c.get("evidence", [])
        ))
        key = (c.get("title"), ev_ids)
        if key in seen_ev:
            continue
        seen_ev.add(key)
        deduped.append(c)

    # Sortiranje po skoru (najjače korelacije prve) — deterministički
    deduped.sort(key=lambda c: (-c.get("score", 0), c.get("id", "")))
    return deduped


# Labele za prikaz pojedinačnih tipova artefakata u narativnoj
# rekonstrukciji događaja ("REKONSTRUKCIJA DOGAĐAJA").
HEADLINE_TYPE_LABELS = {
    "call": "POZIV",
    "location": "LOKACIJA",
    "crypto": "KRIPTO TRANSAKCIJA/ADRESA",
    "comm": "KOMUNIKACIJA",
    "app": "APLIKACIJA",
    "media": "MEDIJA",
    "account": "NALOG/IDENTITET",
    "web": "WEB AKTIVNOST",
}


def _narrate_event(e: dict) -> str:
    """Pretvara jedan headline artefakat u kratku narativnu rečenicu/redak."""
    etype = e.get("type", "?")
    label = HEADLINE_TYPE_LABELS.get(etype, etype.upper())
    value = e.get("value", "")
    source = e.get("source", "")
    line = f"[{label}] {value}"
    if source:
        line += f"  (izvor: {source})"
    return line


def _collect_report_data(session: dict) -> dict:
    """Prikuplja sve podatke potrebne za generisanje izveštaja (text/PDF/DOCX)."""
    results   = session.get("results", {})
    dump_path = session.get("dump_path", "N/A")
    created   = session.get("created_at", "N/A")

    dev = results.get("device_info", {})
    dev_findings = {f["key"]: f["value"] for f in dev.get("findings", [])}
    manuf = dev_findings.get("Proizvođač", "")
    model = dev_findings.get("Model", "")
    if manuf and manuf != "Nepoznat" and model and model != "Nepoznat":
        device_str = f"{manuf} {model}"
    else:
        device_str = model or manuf or "Nepoznat uređaj"
    android_str = dev_findings.get("Android verzija", "")

    all_alerts = []
    for module_name, data in results.items():
        for alert in (data.get("alerts") or []):
            all_alerts.append((module_name, alert))

    timeline = build_timeline(results)
    correlations = build_correlations(results)
    headline = build_headline_timeline(results, correlations)
    total_artifacts = sum(len(data.get("artifacts") or []) for data in results.values())

    return {
        "results": results,
        "dump_path": dump_path,
        "created": created,
        "device_str": device_str,
        "android_str": android_str,
        "all_alerts": all_alerts,
        "timeline": timeline,
        "correlations": correlations,
        "headline": headline,
        "total_artifacts": total_artifacts,
        "now": datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC"),
        # DETERMINISTIČKI otisak SADRŽAJA dokaza (bez vremena/uuid-a). Isti dump
        # → isti heš → dokaz da izveštaj počiva na nepromenjenim nalazima.
        "content_hash": _results_content_hash(results),
        "tool_version": TOOL_VERSION,
    }


def generate_text_report(session: dict) -> str:
    rd = _collect_report_data(session)
    results       = rd["results"]
    dump_path      = rd["dump_path"]
    created        = rd["created"]
    device_str     = rd["device_str"]
    android_str    = rd["android_str"]
    all_alerts     = rd["all_alerts"]
    timeline       = rd["timeline"]
    correlations   = rd["correlations"]
    headline       = rd["headline"]
    total_artifacts = rd["total_artifacts"]
    now            = rd["now"]

    sep = "═" * 68
    thin = "─" * 68

    lines = [
        sep,
        "  FORENZIČKI IZVEŠTAJ — ANDROID MOBILNA FORENZIKA",
        sep,
        "",
        f"  Uređaj:          {device_str}",
        f"  Android:         {android_str}",
        f"  Dump putanja:    {dump_path}",
        f"  Sesija otvorena: {created}",
        f"  Izveštaj generisan: {now}",
        "",
        thin,
        "  METODOLOGIJA",
        thin,
        "  Logički filesystem dump, read-only pristup.",
        "  Alati: Android Forensic Dashboard v1.0",
        "  Integritet: Originalni dump nije modifikovan.",
        "",
    ]

    # ── Izvršni rezime ──────────────────────────────────────────────────
    lines += [
        thin,
        "  IZVRŠNI REZIME",
        thin,
        f"  Analizirano modula:           {len(results)}",
        f"  Ukupno artefakata:             {total_artifacts}",
        f"  Upozorenja:                    {len(all_alerts)}",
        f"  Korelacije između izvora:      {len(correlations)}",
        f"  Događaja u rekonstrukciji:      {len(headline)}",
        "",
    ]

    # ── Rekonstrukcija događaja (headline timeline, narativno) ──────────
    lines += [
        thin,
        f"  REKONSTRUKCIJA DOGAĐAJA ({len(headline)})",
        thin,
        "  Hronološki prikaz najznačajnijih događaja iz svih izvora —",
        "  pozivi, lokacije, kripto transakcije, šifrovana komunikacija,",
        "  sumnjive aplikacije i identitetski artefakti.",
        "",
    ]
    if headline:
        last_date = None
        for e in headline:
            ts = e.get("ts") or ""
            date_part, _, time_part = ts.replace("Z", "").partition("T")
            if date_part != last_date:
                lines.append(f"\n  ── {date_part} ──")
                last_date = date_part
            lines.append(f"  {time_part or '??:??:??'}  {_narrate_event(e)}")
    else:
        lines.append("  Nema dovoljno podataka za rekonstrukciju događaja.")
    lines.append("")

    # ── Upozorenja ────────────────────────────────────────────────────
    lines += [
        thin,
        f"  UPOZORENJA ({len(all_alerts)})",
        thin,
    ]
    for module_name, alert in all_alerts:
        lines.append(f"  [!] [{module_name}] {alert}")
    if not all_alerts:
        lines.append("  Nema upozorenja.")
    lines.append("")

    # ── Rezultati po modulima ────────────────────────────────────────────
    lines += [
        thin,
        "  REZULTATI PO MODULIMA",
        thin,
    ]
    for module_name, data in results.items():
        lines.append(f"\n  ▸ {module_name.upper()}")
        for f in (data.get("findings") or []):
            key_padded = f["key"].ljust(30)
            lines.append(f"    {key_padded} {f['value']}")

    # ── Korelacije ─────────────────────────────────────────────────────
    lines += [
        "",
        thin,
        f"  KORELACIJE ({len(correlations)})",
        thin,
    ]
    for c in correlations:
        lines.append(f"\n  [{c['id']}] {c['title']}")
        lines.append(f"  Pouzdanost: {c['confidence']}")
        lines.append(f"  Izvori: {' + '.join(c['sources'])}")
        lines.append(f"  {c['detail']}")
    if not correlations:
        lines.append("  Nema pronađenih korelacija između izvora.")

    # ── Detaljna vremenska linija ─────────────────────────────────────────
    MAX_DETAILED_TIMELINE = 500
    shown_timeline = timeline[:MAX_DETAILED_TIMELINE]
    lines += [
        "",
        thin,
        f"  DETALJNA VREMENSKA LINIJA ({len(timeline)} događaja)",
        thin,
    ]
    if len(timeline) > MAX_DETAILED_TIMELINE:
        lines.append(
            f"  Prikazano je prvih {MAX_DETAILED_TIMELINE} od {len(timeline)} događaja "
            f"(sortirano hronološki). Kompletna lista je dostupna preko API-ja "
            f"(/api/session/{{id}}/timeline)."
        )
        lines.append("")
    for e in shown_timeline:
        ts = (e.get("ts") or "N/A").replace("T", " ").replace("Z", "")
        lines.append(f"  {ts}  [{e.get('type','?').upper():8}]  {e.get('value','')}")
        lines.append(f"  {'':20}  ◁ {e.get('source','')}  [{e.get('module','')}]")

    lines += [
        "",
        sep,
        "  FORENZIČKI ZAKLJUČAK",
        sep,
        "",
        "  Ovaj izveštaj sadrži faktografske nalaze (objektivni podaci).",
        "  Pravna interpretacija i krivična kvalifikacija isključivo su",
        "  u nadležnosti organa postupka.",
        "",
        f"  Android Forensic Dashboard v1.0 — {now}",
        sep,
    ]

    return "\n".join(lines)


def _build_pdf_report(session: dict) -> bytes:
    """Generiše stilizovan PDF izveštaj (reportlab)."""
    rd = _collect_report_data(session)
    results = rd["results"]

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=2 * cm, bottomMargin=2 * cm, leftMargin=2 * cm, rightMargin=2 * cm,
        title="Forenzicki izvestaj",
    )

    styles = getSampleStyleSheet()
    F = _PDF_FONT
    FB = _PDF_FONT_BOLD
    title_style = ParagraphStyle("TitleX", parent=styles["Title"], fontName=FB, fontSize=18, spaceAfter=4)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontName=FB, fontSize=13, spaceBefore=14, spaceAfter=6,
                         textColor=colors.HexColor("#1f2937"))
    h3 = ParagraphStyle("H3", parent=styles["Heading3"], fontName=FB, fontSize=10, spaceBefore=8, spaceAfter=3,
                         textColor=colors.HexColor("#374151"))
    body = ParagraphStyle("BodyX", parent=styles["Normal"], fontName=F, fontSize=9, leading=13)
    small = ParagraphStyle("Small", parent=styles["Normal"], fontName=F, fontSize=7.5, leading=10)
    muted = ParagraphStyle("Muted", parent=body, fontName=F, fontSize=8, textColor=colors.HexColor("#6b7280"))
    alert_style = ParagraphStyle("Alert", parent=body, fontName=F, textColor=colors.HexColor("#b91c1c"))

    el = []
    el.append(Paragraph("Forenzički izveštaj — Android mobilna forenzika", title_style))
    el.append(Paragraph(f"Generisano: {_esc(rd['now'])}", muted))
    el.append(Spacer(1, 10))

    # ── Osnovni podaci ──
    meta_rows = [
        ["Uređaj", _esc(rd["device_str"])],
        ["Android", _esc(rd["android_str"] or "N/A")],
        ["Dump putanja", _esc(rd["dump_path"])],
        ["Sesija otvorena", _esc(rd["created"])],
    ]
    meta_table = Table(meta_rows, colWidths=[4 * cm, 13 * cm])
    meta_table.setStyle(TableStyle([
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#6b7280")),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("LINEBELOW", (0, 0), (-1, -1), 0.5, colors.HexColor("#e5e7eb")),
    ]))
    el.append(meta_table)
    el.append(Spacer(1, 12))

    # ── Izvršni rezime ──
    el.append(Paragraph("Izvršni rezime", h2))
    summary_rows = [
        ["Analizirano modula", str(len(results))],
        ["Ukupno artefakata", str(rd["total_artifacts"])],
        ["Upozorenja", str(len(rd["all_alerts"]))],
        ["Korelacije između izvora", str(len(rd["correlations"]))],
        ["Događaja u rekonstrukciji", str(len(rd["headline"]))],
    ]
    summary_table = Table(summary_rows, colWidths=[9 * cm, 4 * cm])
    summary_table.setStyle(TableStyle([
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f3f4f6")),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e5e7eb")),
    ]))
    el.append(summary_table)

    # ── Rekonstrukcija događaja ──
    el.append(Paragraph(f"Rekonstrukcija događaja ({len(rd['headline'])})", h2))
    if rd["headline"]:
        last_date = None
        for e in rd["headline"]:
            ts = e.get("ts") or ""
            date_part, _, time_part = ts.replace("Z", "").partition("T")
            if date_part != last_date:
                el.append(Paragraph(date_part, h3))
                last_date = date_part
            label = HEADLINE_TYPE_LABELS.get(e.get("type"), (e.get("type") or "?").upper())
            text = f"<b>{time_part or '??:??:??'}</b> &nbsp; [{_esc(label)}] {_esc(e.get('value', ''))}"
            if e.get("source"):
                text += f" <font color='#9ca3af' size=7>— {_esc(e.get('source'))}</font>"
            el.append(Paragraph(text, body))
    else:
        el.append(Paragraph("Nema dovoljno podataka za rekonstrukciju događaja.", muted))

    # ── Upozorenja ──
    el.append(Paragraph(f"Upozorenja ({len(rd['all_alerts'])})", h2))
    if rd["all_alerts"]:
        for module_name, alert in rd["all_alerts"]:
            el.append(Paragraph(f"⚠ <b>[{_esc(module_name)}]</b> {_esc(alert)}", alert_style))
    else:
        el.append(Paragraph("Nema upozorenja.", muted))

    # ── Rezultati po modulima ──
    el.append(Paragraph("Rezultati po modulima", h2))
    for module_name, data in results.items():
        findings = data.get("findings") or []
        if not findings:
            continue
        el.append(Paragraph(module_name.upper(), h3))
        rows = [[Paragraph(_esc(f["key"]), small), Paragraph(_esc(str(f["value"])), small)] for f in findings]
        t = Table(rows, colWidths=[6 * cm, 11 * cm])
        t.setStyle(TableStyle([
            ("LINEBELOW", (0, 0), (-1, -1), 0.4, colors.HexColor("#f3f4f6")),
            ("TOPPADDING", (0, 0), (-1, -1), 2),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ]))
        el.append(t)

    # ── Korelacije ──
    el.append(Paragraph(f"Korelacije ({len(rd['correlations'])})", h2))
    if rd["correlations"]:
        for c in rd["correlations"]:
            el.append(Paragraph(
                f"[{_esc(c['id'])}] {_esc(c['title'])} &nbsp; "
                f"<font color='#7c3aed'>{_esc(c['confidence'])}</font>", h3))
            el.append(Paragraph(f"Izvori: {_esc(' + '.join(c['sources']))}", muted))
            el.append(Paragraph(_esc(c["detail"]), body))
            el.append(Spacer(1, 4))
    else:
        el.append(Paragraph("Nema pronađenih korelacija između izvora.", muted))

    # ── Detaljna vremenska linija (sažeta i grupisana po danu) ──
    MAX_PDF_TIMELINE = 400
    timeline = rd["timeline"]
    display_rows = collapse_timeline_for_report(timeline)
    raw_web = sum(1 for e in timeline if e.get("type") == "web")
    kept_web = sum(1 for r in display_rows if r.get("kind") == "burst")

    el.append(Paragraph(f"Detaljna vremenska linija ({len(timeline)} događaja)", h2))
    note = (
        f"Uzastopne web posete (browser istorija) su sažete radi preglednosti; "
        f"{raw_web} web događaja prikazano je kroz {kept_web} sažetih blokova. "
        f"Događaji su grupisani po danu i obojeni po ozbiljnosti "
        f"(<font color='#b91c1c'>crveno = visoka</font>, "
        f"<font color='#b45309'>narandžasto = srednja</font>)."
    )
    el.append(Paragraph(note, muted))
    if len(display_rows) > MAX_PDF_TIMELINE:
        el.append(Paragraph(
            f"Prikazano je prvih {MAX_PDF_TIMELINE} od {len(display_rows)} redova; "
            f"kompletna linija je dostupna u aplikaciji.", muted))
    el.append(Spacer(1, 4))

    SEV_BG = {"high": colors.HexColor("#fef2f2"), "medium": colors.HexColor("#fffbeb")}
    SEV_FG = {"high": "#b91c1c", "medium": "#b45309"}

    rows = [["Vreme", "Tip", "Vrednost", "Izvor"]]
    style_cmds = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f2937")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, 0), 8),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#e5e7eb")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
    ]

    last_date = None
    ridx = 0
    for r in display_rows[:MAX_PDF_TIMELINE]:
        ts = r.get("ts") or ""
        date_part, _, time_part = ts.replace("Z", "").partition("T")
        # Dan-separator red preko cele širine
        if date_part and date_part != last_date:
            last_date = date_part
            ridx += 1
            rows.append([Paragraph(f"<b>{_esc(date_part)}</b>", small), "", "", ""])
            style_cmds += [
                ("SPAN", (0, ridx), (-1, ridx)),
                ("BACKGROUND", (0, ridx), (-1, ridx), colors.HexColor("#eef2ff")),
                ("TEXTCOLOR", (0, ridx), (-1, ridx), colors.HexColor("#3730a3")),
            ]

        # Vreme (za burst prikaži opseg)
        t_disp = (time_part or "")[:8]
        if r.get("kind") == "burst" and r.get("ts_end"):
            _, _, tend = (r["ts_end"] or "").replace("Z", "").partition("T")
            t_disp = f"{t_disp}–{tend[:8]}"

        sev = r.get("severity", "low")
        val_txt = _esc(r.get("value", ""))
        if sev in SEV_FG:
            val_txt = f"<font color='{SEV_FG[sev]}'>{val_txt}</font>"

        ridx += 1
        rows.append([
            Paragraph(_esc(t_disp), small),
            Paragraph(_esc((r.get("type") or "?").upper()), small),
            Paragraph(val_txt, small),
            Paragraph(_esc(r.get("source", "")), small),
        ])
        if sev in SEV_BG:
            style_cmds.append(("BACKGROUND", (0, ridx), (-1, ridx), SEV_BG[sev]))

    t = Table(rows, colWidths=[2.6 * cm, 1.8 * cm, 8.6 * cm, 4 * cm], repeatRows=1)
    t.setStyle(TableStyle(style_cmds))
    el.append(t)

    el.append(Spacer(1, 14))
    el.append(HRFlowable(width="100%", color=colors.HexColor("#e5e7eb")))
    el.append(Paragraph(
        "Ovaj izveštaj sadrži faktografske nalaze (objektivni podaci). Pravna interpretacija "
        "i krivična kvalifikacija isključivo su u nadležnosti organa postupka.", muted))

    doc.build(el)
    return buf.getvalue()


def _build_docx_report(session: dict) -> bytes:
    """Generiše Word (.docx) izveštaj sa istom strukturom kao PDF/text verzija."""
    rd = _collect_report_data(session)
    results = rd["results"]

    doc = Document()
    doc.add_heading("Forenzički izveštaj — Android mobilna forenzika", level=0)
    doc.add_paragraph(f"Generisano: {rd['now']}")

    doc.add_heading("Osnovni podaci", level=1)
    table = doc.add_table(rows=0, cols=2)
    for k, v in [
        ("Uređaj", rd["device_str"]),
        ("Android", rd["android_str"] or "N/A"),
        ("Dump putanja", rd["dump_path"]),
        ("Sesija otvorena", rd["created"]),
    ]:
        row = table.add_row().cells
        row[0].text = k
        row[1].text = str(v)

    doc.add_heading("Izvršni rezime", level=1)
    t2 = doc.add_table(rows=0, cols=2)
    for k, v in [
        ("Analizirano modula", len(results)),
        ("Ukupno artefakata", rd["total_artifacts"]),
        ("Upozorenja", len(rd["all_alerts"])),
        ("Korelacije između izvora", len(rd["correlations"])),
        ("Događaja u rekonstrukciji", len(rd["headline"])),
    ]:
        row = t2.add_row().cells
        row[0].text = k
        row[1].text = str(v)

    doc.add_heading(f"Rekonstrukcija događaja ({len(rd['headline'])})", level=1)
    if rd["headline"]:
        last_date = None
        for e in rd["headline"]:
            ts = e.get("ts") or ""
            date_part, _, time_part = ts.replace("Z", "").partition("T")
            if date_part != last_date:
                doc.add_heading(date_part, level=2)
                last_date = date_part
            label = HEADLINE_TYPE_LABELS.get(e.get("type"), (e.get("type") or "?").upper())
            p = doc.add_paragraph()
            p.add_run(f"{time_part or '??:??:??'}  ").bold = True
            p.add_run(f"[{label}] {e.get('value', '')}")
            if e.get("source"):
                run = p.add_run(f"  — {e.get('source')}")
                run.italic = True
    else:
        doc.add_paragraph("Nema dovoljno podataka za rekonstrukciju događaja.")

    doc.add_heading(f"Upozorenja ({len(rd['all_alerts'])})", level=1)
    if rd["all_alerts"]:
        for module_name, alert in rd["all_alerts"]:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"[{module_name}] ").bold = True
            p.add_run(alert)
    else:
        doc.add_paragraph("Nema upozorenja.")

    doc.add_heading("Rezultati po modulima", level=1)
    for module_name, data in results.items():
        findings = data.get("findings") or []
        if not findings:
            continue
        doc.add_heading(module_name.upper(), level=2)
        t = doc.add_table(rows=0, cols=2)
        for f in findings:
            row = t.add_row().cells
            row[0].text = f["key"]
            row[1].text = str(f["value"])

    doc.add_heading(f"Korelacije ({len(rd['correlations'])})", level=1)
    if rd["correlations"]:
        for c in rd["correlations"]:
            doc.add_heading(f"[{c['id']}] {c['title']}  —  {c['confidence']}", level=2)
            p = doc.add_paragraph()
            p.add_run("Izvori: ").bold = True
            p.add_run(" + ".join(c["sources"]))
            doc.add_paragraph(c["detail"])
    else:
        doc.add_paragraph("Nema pronađenih korelacija između izvora.")

    MAX_DOCX_TIMELINE = 300
    timeline = rd["timeline"]
    doc.add_heading(f"Detaljna vremenska linija ({len(timeline)} događaja)", level=1)
    if len(timeline) > MAX_DOCX_TIMELINE:
        doc.add_paragraph(
            f"Prikazano je prvih {MAX_DOCX_TIMELINE} od {len(timeline)} događaja "
            f"(sortirano hronološki). Kompletna lista je dostupna u aplikaciji."
        )
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Vreme", "Tip", "Vrednost", "Izvor"
    for e in timeline[:MAX_DOCX_TIMELINE]:
        ts = (e.get("ts") or "N/A").replace("T", " ").replace("Z", "")
        row = t.add_row().cells
        row[0].text = ts
        row[1].text = (e.get("type") or "?").upper()
        row[2].text = e.get("value", "")
        row[3].text = e.get("source", "")

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "Ovaj izveštaj sadrži faktografske nalaze (objektivni podaci). Pravna interpretacija "
        "i krivična kvalifikacija isključivo su u nadležnosti organa postupka."
    )
    run.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── Endpointi ────────────────────────────────────────────────────────────

@app.post("/api/session")
def create_session(body: OpenDumpRequest):
    dump_path = body.dump_path.strip()
    if not Path(dump_path).exists():
        raise HTTPException(status_code=400, detail=f"Putanja ne postoji: {dump_path}")

    resolver = DumpResolver(dump_path)
    summary = resolver.summary()

    # Izvuci osnovne info o uređaju za UI
    device_label = "Android uređaj"
    android_label = ""
    build_prop = resolver.resolve("build_prop")
    if build_prop:
        try:
            props = {}
            for line in build_prop.read_text(errors="replace").splitlines():
                if "=" in line and not line.startswith("#"):
                    k, _, v = line.partition("=")
                    props[k.strip()] = v.strip()
            model = props.get("ro.product.model", "")
            mfr   = props.get("ro.product.manufacturer", "")
            ver   = props.get("ro.build.version.release", "")
            sdk   = props.get("ro.build.version.sdk", "")
            if model:
                device_label = f"{mfr} {model}".strip()
            if ver:
                android_label = f"Android {ver} (SDK {sdk})"
        except Exception:
            pass

    session_id = str(uuid.uuid4())

    # Perzistentni slučaj (chain of custody) — ne gubi se na restart.
    # Ako dump dolazi iz akvizicije (fs_case_id), taj Case_YYYY_NNNN se upisuje kao
    # case_number → analitički slučaj i akvizicioni slučaj su povezani.
    case = case_store.create_case(case_number=body.fs_case_id or "", title=device_label,
                                  examiner=body.examiner or "nepoznat")
    audit_log.log_event(
        actor=f"examiner:{body.examiner or 'nepoznat'}", action="open_case",
        case_id=case["case_id"], params={"dump_path": dump_path, "device": device_label,
                                         "source": body.source, "fs_case_id": body.fs_case_id})

    SESSIONS[session_id] = {
        "dump_path": dump_path,
        "resolver": resolver,
        "results": {},
        "created_at": datetime.now(timezone.utc).isoformat(),
        "case_id": case["case_id"],
        "fs_case_id": body.fs_case_id or "",
        "source": body.source or "dump",
        "examiner": body.examiner or "",
    }

    logger.info(f"Session {session_id[:8]} opened: {dump_path} (case {case['case_id']})")

    return {
        "session_id": session_id,
        "case_id": case["case_id"],
        "summary": {
            "dump_path": dump_path,
            "device": device_label,
            "android": android_label,
            "found_artifacts": list(summary.get("found_artifacts", {}).keys()),
            "installed_packages": summary.get("installed_packages", [])[:20],
        },
    }


@app.get("/api/session/{session_id}")
def get_session_info(session_id: str):
    session = get_session(session_id)
    return {
        "session_id": session_id,
        "dump_path": session["dump_path"],
        "created_at": session["created_at"],
        "analyzed_modules": list(session["results"].keys()),
    }


@app.delete("/api/session/{session_id}")
def delete_session(session_id: str):
    session = get_session(session_id)
    # Soft-archive slučaja umesto tihog uništenja + audit trag (chain of custody).
    case_id = session.get("case_id")
    if case_id:
        case_store.archive_case(case_id)
        audit_log.log_event(actor="system", action="close_session", case_id=case_id,
                            params={"note": "session zatvorena, slučaj arhiviran (podaci zadržani)"})
    del SESSIONS[session_id]
    return {"deleted": session_id, "case_archived": case_id}


@app.get("/api/cases")
def list_all_cases():
    """Svi slučajevi (multi-case), perzistentni preko restarta."""
    return {"cases": case_store.list_cases()}


@app.get("/api/correlations/rules")
def get_correlation_rules():
    """Deklarativni registry korelacionih pravila (auditabilnost, #4)."""
    return correlation_registry()


@app.get("/api/session/{session_id}/media")
def get_media(session_id: str, kind: str | None = None, only_gps: bool = False,
              album: str | None = None, limit: int = 2000):
    """
    Galerija: lista SVIH slika i snimaka u korisničkim media direktorijumima,
    obogaćena metapodacima iz EXIF modula (GPS, vreme, uređaj, stego, hash).
    `rel` se koristi kao path za /file endpoint (thumbnail + pun pregled).
    """
    from modules import exif as exif_mod
    session = get_session(session_id)
    resolver: DumpResolver = session["resolver"]
    root = resolver.root.resolve()

    # Indeksiraj EXIF artefakte po imenu fajla radi metapodataka
    meta_by_name = {}
    for a in (session["results"].get("exif", {}).get("artifacts") or []):
        ex = a.get("extra") or {}
        fn = ex.get("filename")
        if fn and fn not in meta_by_name:
            meta_by_name[fn] = a

    IMG, VID = exif_mod.IMAGE_EXTENSIONS, exif_mod.VIDEO_EXTENSIONS
    seen, items = set(), []
    for rel_dir in exif_mod.SEARCH_DIRS:
        d = root / rel_dir
        if not d.exists():
            continue
        for f in sorted(d.rglob("*")):
            if not f.is_file():
                continue
            ext = f.suffix.lower()
            k = "image" if ext in IMG else ("video" if ext in VID else None)
            if not k or (kind and kind != k):
                continue
            try:
                rel = str(f.relative_to(root)).replace("\\", "/")  # URL-safe, cross-platform
            except Exception:
                continue
            if rel in seen:
                continue
            seen.add(rel)

            a = meta_by_name.get(f.name)
            ex = (a.get("extra") if a else {}) or {}
            has_gps = ex.get("lat") is not None and ex.get("lon") is not None
            if only_gps and not has_gps:
                continue
            # Album = naziv roditeljskog foldera (Camera, Screenshots, Instagram,
            # WhatsApp Images...) — za grupisanje galerije po albumima.
            parts = rel.split("/")
            album_name = parts[-2] if len(parts) >= 2 else "Ostalo"
            if album_name in (".thumbnails", "thumbnails"):
                album_name = "Thumbnails"
            if album and album != album_name:
                continue
            items.append({
                "filename": f.name,
                "rel": rel,
                "album": album_name,
                "kind": k,
                "size_kb": (f.stat().st_size // 1024) if f.exists() else 0,
                "ts": (a.get("ts") if a else None),
                "lat": ex.get("lat"), "lon": ex.get("lon"),
                "device": ex.get("device") or ex.get("make"),
                "stego": bool(ex.get("stego")),
                "sha256": (a.get("hash_set") or {}).get("sha256") if a else None,
                "artifact_id": a.get("id") if a else None,
            })
            if len(items) >= limit:
                break
        if len(items) >= limit:
            break

    # GPS/sumnjive prve, pa po vremenu
    items.sort(key=lambda m: (0 if (m["lat"] is not None or m["stego"]) else 1, m["ts"] or ""))

    # Rezime po albumima (za grupisanje u galeriji)
    album_counts = {}
    for m in items:
        album_counts[m["album"]] = album_counts.get(m["album"], 0) + 1
    albums = sorted(({"name": n, "count": c} for n, c in album_counts.items()),
                    key=lambda x: -x["count"])

    return {"count": len(items),
            "with_gps": sum(1 for m in items if m["lat"] is not None),
            "stego": sum(1 for m in items if m["stego"]),
            "albums": albums,
            "media": items}


@app.get("/api/session/{session_id}/case")
def get_case_info(session_id: str):
    """Detalji slučaja: runs (verzije), audit trag, provera reproducibilnosti."""
    session = get_session(session_id)
    case_id = session.get("case_id")
    if not case_id:
        raise HTTPException(status_code=404, detail="Slučaj nije vezan za sesiju")
    runs = case_store.list_runs(case_id)
    return {
        "case": case_store.get_case(case_id),
        "runs": runs,
        "audit": audit_log.read_events(case_id, limit=200),
        "audit_chain": audit_log.verify_chain(),
        "reproducible": _reproducibility_check(runs),
    }


def _reproducibility_check(runs: list) -> dict:
    """Da li uzastopni run-ovi istog dump-a daju isti result_hash (determinizam)."""
    hashes = [r.get("result_hash") for r in runs if r.get("result_hash")]
    if len(hashes) < 2:
        return {"verifiable": False, "reason": "Potrebna 2+ analize istog dump-a"}
    return {"verifiable": True, "reproducible": len(set(hashes)) == 1,
            "runs_compared": len(hashes), "distinct_hashes": len(set(hashes))}


@app.get("/api/session/{session_id}/analyze/{module_name}")
def analyze_module(session_id: str, module_name: str):
    session = get_session(session_id)

    if module_name not in MODULE_MAP:
        # Vrati "not implemented yet" za module koji još nisu napisani
        result = {
            "status": "not_found",
            "findings": [{"key": "Status", "value": f"Modul '{module_name}' još nije implementiran"}],
            "artifacts": [],
            "alerts": [],
        }
        session["results"][module_name] = result
        return result

    try:
        logger.info(f"Session {session_id[:8]} running module: {module_name}")
        result = MODULE_MAP[module_name](session["dump_path"])
        # Normalization layer: obogati artefakte (hash, provenance, id, confidence)
        result = _enrich_module(session, module_name, result)
        session["results"][module_name] = result
        audit_log.log_event(actor="system", action="analyze_module",
                            case_id=session.get("case_id"),
                            params={"module": module_name, "artifacts": len(result.get("artifacts") or [])})
        return result
    except Exception as e:
        logger.error(f"Module {module_name} error: {e}", exc_info=True)
        result = {
            "status": "error",
            "findings": [{"key": "Greška", "value": str(e)}],
            "artifacts": [],
            "alerts": [f"Greška u modulu {module_name}: {str(e)}"],
            "error": str(e),
        }
        session["results"][module_name] = result
        return result


@app.post("/api/session/{session_id}/analyze/all")
async def analyze_all(session_id: str):
    """
    Pokreni sve module KONKURENTNO (thread pool) umesto sekvencijalno, obogati
    artefakte (normalization layer), perzistuj kao immutable run i upiši audit.
    """
    session = get_session(session_id)
    dump_path = session["dump_path"]
    run_id = None

    def _run_one(module_name, fn):
        try:
            result = fn(dump_path)
            return module_name, _enrich_module(session, module_name, result)
        except Exception as e:
            logger.error(f"Module {module_name} error: {e}")
            case_id = session.get("case_id")
            # strukturisano praćenje grešaka (perzistentno)
            try:
                if session.get("_run_id"):
                    case_store.save_error(session["_run_id"], module_name, type(e).__name__, str(e), traceback.format_exc())
            except Exception:
                pass
            return module_name, {"status": "error", "findings": [], "artifacts": [],
                                 "alerts": [str(e)], "error": str(e)}

    # Konkurentno izvršavanje u bounded thread pool-u; ne blokira event loop.
    loop = asyncio.get_event_loop()
    with ThreadPoolExecutor(max_workers=min(6, len(MODULE_MAP))) as pool:
        tasks = [loop.run_in_executor(pool, _run_one, name, fn) for name, fn in MODULE_MAP.items()]
        completed = await asyncio.gather(*tasks)

    # Deterministički redosled (po MODULE_MAP), ne po redosledu završetka
    order = list(MODULE_MAP.keys())
    results = {name: res for name, res in sorted(completed, key=lambda kv: order.index(kv[0]))}
    session["results"] = results

    # Perzistuj kao immutable run + audit
    case_id = session.get("case_id")
    if case_id:
        result_hash = _results_content_hash(results)
        run_id = case_store.save_run(case_id, dump_path, results,
                                     input_fingerprint=_dump_fingerprint(session["resolver"]),
                                     tool_version=TOOL_VERSION, result_hash=result_hash)
        session["_run_id"] = run_id
        audit_log.log_event(actor="system", action="analyze_all", case_id=case_id, run_id=run_id,
                            params={"modules": len(results), "result_hash": result_hash})

    return {"results": results, "run_id": run_id}


@app.get("/api/session/{session_id}/correlations")
def get_correlations(session_id: str):
    session = get_session(session_id)
    return build_correlations(session["results"])


@app.get("/api/session/{session_id}/correlations/graph")
def get_correlation_graph(session_id: str):
    """
    Entity-relationship graf za vizualizaciju: čvorovi su forenzički
    entiteti (telefoni, emailovi, kripto adrese, paketi, SSID...),
    ivice su ko-pojavljivanja i korelacione veze.
    """
    session = get_session(session_id)
    results = session["results"]
    correlations = build_correlations(results)
    return build_entity_graph(results, correlations)


@app.get("/api/session/{session_id}/dashboard")
def get_dashboard(session_id: str):
    """Izvršni rezime za dashboard + status otkrivanja artefakata."""
    session = get_session(session_id)
    results = session["results"]
    resolver: DumpResolver = session["resolver"]

    correlations = build_correlations(results)
    timeline = build_timeline(results)
    headline = build_headline_timeline(results, correlations)
    all_alerts = [a for data in results.values() for a in (data.get("alerts") or [])]

    severity_breakdown = {"high": 0, "medium": 0, "low": 0}
    for e in timeline:
        severity_breakdown[e.get("severity", "low")] = severity_breakdown.get(e.get("severity", "low"), 0) + 1

    return {
        "analyzed_modules": len(results),
        "total_artifacts": sum(len(d.get("artifacts") or []) for d in results.values()),
        "alerts": len(all_alerts),
        "correlations": len(correlations),
        "timeline_events": len(timeline),
        "headline_events": len(headline),
        "severity_breakdown": severity_breakdown,
        # Temporalne anomalije + praznine u vremenskoj liniji (#3).
        "timeline_anomalies": detect_timeline_anomalies(timeline),
        # Numerička raspodela pouzdanosti artefakata (0-100).
        "confidence_distribution": _confidence_distribution(results),
        # Kako je svaka ključna baza pronađena (kanonski / regex / šema / nije) —
        # radi za bilo koji uređaj, ne samo AOSP Android 10.
        "artifact_discovery": resolver.discovery_report(),
        # Rezime popisa SVIH baza u dump-u (koliko, koliko prepoznato).
        "database_inventory": inventory_summary(scan_all_databases(resolver)),
        "ai_available": ai_available()[0],
        "case_id": session.get("case_id"),
        "reproducible_hash": _results_content_hash(results),
    }


def _confidence_distribution(results: dict) -> dict:
    """Raspodela numeričke pouzdanosti artefakata u opsege (0-25/26-50/51-75/76-100)."""
    buckets = {"0-25": 0, "26-50": 0, "51-75": 0, "76-100": 0, "bez": 0}
    for data in results.values():
        for a in (data.get("artifacts") or []):
            c = a.get("confidence")
            if c is None:
                buckets["bez"] += 1
            elif c <= 25:
                buckets["0-25"] += 1
            elif c <= 50:
                buckets["26-50"] += 1
            elif c <= 75:
                buckets["51-75"] += 1
            else:
                buckets["76-100"] += 1
    return buckets


@app.get("/api/session/{session_id}/databases")
def get_databases(session_id: str):
    """
    Popis SVIH SQLite baza u dump-u sa automatskom klasifikacijom (po
    tabelama i imenu). Ništa se ne 'gubi' — čak i neprepoznate baze se
    prijavljuju sa tabelama kao tragom. Radi za bilo koji uređaj/aplikaciju.
    """
    session = get_session(session_id)
    resolver: DumpResolver = session["resolver"]
    inventory = scan_all_databases(resolver)
    return {
        "summary": inventory_summary(inventory),
        "databases": inventory,
    }


@app.get("/api/session/{session_id}/ai-conclusion")
def get_ai_conclusion(session_id: str):
    """
    AI forenzički zaključak — Claude model uzima sve nalaze + popis baza i
    generiše koherentan forenzički zaključak (rekonstrukcija, značaj,
    preporuke). Zahteva ANTHROPIC_API_KEY; ako nema, vraća jasan razlog.
    """
    session = get_session(session_id)
    data = _collect_report_data(session)
    data["db_inventory"] = scan_all_databases(session["resolver"])
    result = generate_ai_conclusion(data)
    audit_log.log_event(actor="system", action="ai_conclusion", case_id=session.get("case_id"),
                        run_id=session.get("_run_id"),
                        params={"available": result.get("available"), "model": result.get("model")})
    return result


@app.get("/api/session/{session_id}/timeline")
def get_timeline(
    session_id: str,
    type: str | None = None,
    search: str | None = None,
    severity: str | None = None,
    limit: int | None = None,
    offset: int = 0,
):
    """
    Hronološka linija sa server-side filtriranjem:
    ?type=location&search=samsung&severity=high&limit=100&offset=0
    Bez parametara vraća kompletnu listu (kompatibilno sa postojećim frontendom).
    """
    session = get_session(session_id)
    events = build_timeline(session["results"])

    if type:
        events = [e for e in events if e.get("type") == type]
    if severity:
        events = [e for e in events if e.get("severity") == severity]
    if search:
        q = search.lower()
        events = [
            e for e in events
            if q in str(e.get("value", "")).lower()
            or q in str(e.get("source", "")).lower()
            or q in str(e.get("module", "")).lower()
        ]
    if offset:
        events = events[offset:]
    if limit is not None:
        events = events[:limit]
    return events


@app.get("/api/session/{session_id}/timeline/headline")
def get_headline_timeline(session_id: str):
    session = get_session(session_id)
    correlations = build_correlations(session["results"])
    return build_headline_timeline(session["results"], correlations)


@app.get("/api/session/{session_id}/report")
def get_report(session_id: str, format: str = "text"):
    session = get_session(session_id)

    if format == "json":
        # Sirovi podaci (kompatibilnost sa postojećim frontendom)
        correlations = build_correlations(session["results"])
        return {
            "results": session["results"],
            "correlations": correlations,
            "timeline": build_timeline(session["results"]),
            "headline_timeline": build_headline_timeline(session["results"], correlations),
        }

    if format == "model":
        # DETERMINISTIČKI kanonski model izveštaja (#8): iz njega se renderuje,
        # AI je samo narativni sloj. Content_hash je reproducibilan.
        audit_log.log_event(actor="system", action="generate_report",
                            case_id=session.get("case_id"), run_id=session.get("_run_id"),
                            params={"format": "model"})
        return build_report_model(_collect_report_data(session))

    date_tag = datetime.now(timezone.utc).strftime("%Y%m%d")

    if format == "pdf":
        pdf_bytes = _build_pdf_report(session)
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="forenzicki_izvestaj_{date_tag}.pdf"'},
        )

    if format == "docx":
        docx_bytes = _build_docx_report(session)
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="forenzicki_izvestaj_{date_tag}.docx"'},
        )

    if format == "html":
        html = ForensicReportEngine(_collect_report_data(session)).render()
        return Response(
            content=html,
            media_type="text/html; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="forenzicki_izvestaj_{date_tag}.html"'},
        )

    return generate_text_report(session)


# ─── Pristup originalnim fajlovima u dump-u (pregled/eksplorer) ────────────

@app.get("/api/session/{session_id}/file")
def get_artifact_file(session_id: str, path: str):
    """
    Servira originalni fajl iz dump-a (za pregled slika/audio/teksta u UI).
    `path` je relativna putanja kao u 'source' polju artefakta.
    """
    session = get_session(session_id)
    file_path = resolve_artifact_file(session, path)
    return FileResponse(file_path)


@app.post("/api/session/{session_id}/reveal")
def reveal_artifact_file(session_id: str, body: RevealRequest):
    """
    Otvara sistemski fajl eksplorer sa selektovanim fajlom (samo lokalno,
    desktop korišćenje aplikacije).
    """
    session = get_session(session_id)
    file_path = resolve_artifact_file(session, body.path)

    try:
        if sys.platform.startswith("win"):
            subprocess.Popen(["explorer", "/select,", str(file_path)])
        elif sys.platform == "darwin":
            subprocess.Popen(["open", "-R", str(file_path)])
        else:
            subprocess.Popen(["xdg-open", str(file_path.parent)])
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Nije moguće otvoriti eksplorer: {e}")

    # Chain of custody: veštak je namerno pristupio dokaznom fajlu — audituj.
    audit_log.log_event(actor="examiner", action="reveal_evidence_file",
                        case_id=session.get("case_id"), run_id=session.get("_run_id"),
                        params={"path": body.path})
    return {"opened": str(file_path)}


# ════════════════════════════════════════════════════════════════════════════
# ACQUISITION LAYER — detekcija izvora + pokretanje akvizicije (job-ovi)
# Analitički engine ostaje netaknut: akvizicija samo napravi Evidence/ folder,
# a onda se poziva postojeći POST /api/session sa tom putanjom.
# ════════════════════════════════════════════════════════════════════════════

@app.get("/api/sources")
def get_sources():
    """Pregled dostupnosti svih izvora (za acquisition wizard)."""
    return acq_detect.sources_overview()


@app.get("/api/detect/phone")
def detect_phone():
    return acq_detect.detect_phones()


@app.get("/api/detect/sim")
def detect_sim():
    return acq_detect.detect_sim_readers()


@app.get("/api/detect/storage")
def detect_storage(kind: str = "sdcard"):
    return acq_detect.detect_storage(kind if kind in ("sdcard", "usb") else "sdcard")


def _acquire_target(source: str):
    """Mapiranje izvora → driver target funkcija (phone/sim se lenjivo uvoze)."""
    if source in ("sdcard", "usb"):
        return acq_storage.acquire_storage
    if source == "mobile":
        from acquisition import phone as acq_phone
        return acq_phone.acquire_phone
    if source == "sim":
        from acquisition import sim as acq_sim
        return acq_sim.acquire_sim
    raise HTTPException(status_code=400, detail=f"Nepoznat izvor akvizicije: {source}")


@app.post("/api/acquire/{source}")
def start_acquisition(source: str, body: AcquireRequest):
    """
    Pokreni akviziciju (asinhrono, u pozadinskoj niti). Vraća job_id za praćenje.
    source: mobile | sim | sdcard | usb
    """
    if source not in ("mobile", "sim", "sdcard", "usb"):
        raise HTTPException(status_code=400, detail=f"Nepoznat izvor: {source}")
    target = _acquire_target(source)

    # Kwargs po tipu izvora (prosleđuju se driver funkciji).
    if source in ("sdcard", "usb"):
        if not body.mount:
            raise HTTPException(status_code=400, detail="Nije izabran disk (mount).")
        kwargs = {"mount": body.mount, "kind": source, "examiner": body.examiner,
                  "disk_info": body.disk_info or {}}
    elif source == "mobile":
        kwargs = {"serial": body.serial, "examiner": body.examiner,
                  "device_info": body.device_info or {}}
    else:  # sim
        kwargs = {"reader_name": body.reader, "examiner": body.examiner}

    job_id = acq_jobs.start_job(source, target, **kwargs)
    audit_log.log_event(actor=f"examiner:{body.examiner or 'nepoznat'}",
                        action="start_acquisition",
                        params={"source": source, "job_id": job_id,
                                "target": body.mount or body.serial or body.reader})
    return {"job_id": job_id, "source": source}


@app.get("/api/acquire/job/{job_id}")
def acquire_job_status(job_id: str):
    j = acq_jobs.get_job(job_id)
    if not j:
        raise HTTPException(status_code=404, detail="Posao nije pronađen.")
    return j


@app.post("/api/acquire/job/{job_id}/cancel")
def acquire_job_cancel(job_id: str):
    ok = acq_jobs.cancel_job(job_id)
    if not ok:
        raise HTTPException(status_code=404, detail="Posao nije pronađen ili nije aktivan.")
    audit_log.log_event(actor="examiner", action="cancel_acquisition", params={"job_id": job_id})
    return {"cancelled": job_id}


@app.get("/api/acquire/cases")
def list_acquisition_cases():
    """Svi slučajevi na disku (central case manager, akvizicija)."""
    return {"cases": acq_cases.list_fs_cases()}


def _acq_report_model(case_id: str) -> dict:
    """Nađi report_data slučaja (iz job rezultata ili case.json) i napravi model."""
    # 1) probaj iz aktivnih job-ova
    for j in acq_jobs.list_jobs():
        res = None
        full = acq_jobs.get_job(j["id"])
        if full and full.get("result") and full["result"].get("case_id") == case_id:
            res = full["result"]
        if res and res.get("report_data"):
            return packager_mod.model_from_acquisition(res["report_data"])
    # 2) fallback: iz case.json (osnovni podaci)
    meta = acq_cases.read_case_meta(case_id)
    if not meta:
        raise HTTPException(status_code=404, detail="Slučaj nije pronađen.")
    rd = {"kind": meta.get("source", "sdcard"), "case_id": case_id,
          "device": meta.get("device_info") or {}, "stats": {},
          "manifest_summary": meta.get("hashes") or {}}
    return packager_mod.model_from_acquisition(rd)


@app.get("/api/acquire/case/{case_id}/report")
def acquisition_report(case_id: str, format: str = "pdf"):
    """Namenski izveštaj akvizicije (SIM/SD/USB) u traženom formatu; upiše i u Reports/."""
    model = _acq_report_model(case_id)
    try:
        packager_mod.write_report_set(case_id, model)  # Reports/Full_Report.{pdf,docx,html,txt}
    except Exception:
        pass
    content, media, ext = exporters_mod.render(model, format)
    audit_log.log_event(actor="system", action="acquisition_report",
                        params={"case_id": case_id, "format": ext})
    return _file_response(content, media, f"{case_id}_izvestaj.{ext}")


@app.get("/api/acquire/case/{case_id}/download")
def acquisition_download(case_id: str, format: str = "zip"):
    """Preuzmi ceo slučaj kao .zip (ili .tar.gz) — Evidence/Reports/Logs/Exports."""
    try:
        if format == "tar":
            data, fname = packager_mod.build_case_tar(case_id)
            media = "application/gzip"
        else:
            data, fname = packager_mod.build_case_zip(case_id)
            media = "application/zip"
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))
    audit_log.log_event(actor="examiner", action="export_case_package",
                        params={"case_id": case_id, "format": format})
    return _file_response(data, media, fname)


# ════════════════════════════════════════════════════════════════════════════
# UNIVERSAL EXPORT — izvezi BILO KOJI prikaz (PDF/DOCX/HTML/TXT) ili ceo slučaj
# ════════════════════════════════════════════════════════════════════════════

def _file_response(content, media_type: str, filename: str) -> Response:
    body = content if isinstance(content, (bytes, bytearray)) else str(content).encode("utf-8")
    return Response(content=body, media_type=media_type,
                    headers={"Content-Disposition": f'attachment; filename="{filename}"'})


def _session_case_meta(session: dict) -> dict:
    results = session.get("results", {})
    dev = results.get("device_info", {})
    dev_findings = {f["key"]: f["value"] for f in dev.get("findings", [])}
    device = (f'{dev_findings.get("Proizvođač","")} {dev_findings.get("Model","")}').strip() or None
    return {
        "case_id": session.get("case_id"),
        "case_number": session.get("fs_case_id"),
        "examiner": session.get("examiner"),
        "source": session.get("source", "dump"),
        "device": device,
        "dump_path": session.get("dump_path"),
        "created_at": session.get("created_at"),
    }


def _flatten_evidence(results: dict) -> list:
    out = []
    for module_name, data in results.items():
        for a in (data.get("artifacts") or []):
            out.append({**a, "module": a.get("module") or module_name})
    return out


def _build_export_model(session: dict, view: str) -> dict:
    results = session.get("results", {})
    meta = _session_case_meta(session)
    if view == "timeline":
        return exporters_mod.model_from_timeline(build_timeline(results), meta)
    if view == "correlations":
        return exporters_mod.model_from_correlations(build_correlations(results), meta)
    if view == "evidence":
        return exporters_mod.model_from_evidence(_flatten_evidence(results), meta)
    if view == "dashboard":
        correlations = build_correlations(results)
        timeline = build_timeline(results)
        headline = build_headline_timeline(results, correlations)
        alerts = [a for d in results.values() for a in (d.get("alerts") or [])]
        pairs = [
            {"label": "Analizirano modula", "value": len(results)},
            {"label": "Ukupno artefakata", "value": sum(len(d.get("artifacts") or []) for d in results.values())},
            {"label": "Upozorenja", "value": len(alerts)},
            {"label": "Korelacije", "value": len(correlations)},
            {"label": "Događaja (rekonstrukcija)", "value": len(headline)},
            {"label": "Događaja (timeline)", "value": len(timeline)},
        ]
        return {"title": "Izvršni rezime (Pregled)", "meta": exporters_mod._case_meta(meta),
                "sections": [{"heading": "Statistika", "type": "keyvalue", "pairs": pairs},
                             {"heading": f"Upozorenja ({len(alerts)})", "type": "list",
                              "items": alerts or ["Nema upozorenja."]}]}
    if view.startswith("module:"):
        mod = view.split(":", 1)[1]
        return exporters_mod.model_from_artifacts(mod, results.get(mod, {}), meta)
    raise HTTPException(status_code=400, detail=f"Nepoznat prikaz za izvoz: {view}")


@app.get("/api/session/{session_id}/export")
def export_view(session_id: str, view: str = "evidence", format: str = "pdf"):
    """
    Univerzalni izvoz prikaza. view: dashboard|timeline|correlations|evidence|module:<id>|report.
    Za 'report' koristi se postojeći puni izveštaj (identično /report endpointu).
    """
    session = get_session(session_id)
    fmt = (format or "pdf").lower()

    if view == "report":
        # Preusmeri na postojeći, bogati izveštaj (ne dupliramo logiku)
        if fmt == "pdf":
            return _file_response(_build_pdf_report(session), "application/pdf",
                                  f"forenzicki_izvestaj.pdf")
        if fmt == "docx":
            return _file_response(_build_docx_report(session),
                                  "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                  "forenzicki_izvestaj.docx")
        if fmt == "html":
            return _file_response(ForensicReportEngine(_collect_report_data(session)).render(),
                                  "text/html; charset=utf-8", "forenzicki_izvestaj.html")
        return _file_response(generate_text_report(session), "text/plain; charset=utf-8",
                              "forenzicki_izvestaj.txt")

    model = _build_export_model(session, view)
    content, media, ext = exporters_mod.render(model, fmt)
    audit_log.log_event(actor="examiner", action="export_view",
                        case_id=session.get("case_id"), run_id=session.get("_run_id"),
                        params={"view": view, "format": ext})
    safe_view = view.replace(":", "_")
    return _file_response(content, media, f"{safe_view}.{ext}")


@app.post("/api/session/{session_id}/export/artifact")
def export_artifact(session_id: str, body: ArtifactExportRequest, format: str = "pdf"):
    """Izvezi pojedinačni artefakt (detaljan izveštaj) u traženom formatu."""
    session = get_session(session_id)
    model = exporters_mod.model_from_single_artifact(body.artifact, _session_case_meta(session))
    content, media, ext = exporters_mod.render(model, format)
    audit_log.log_event(actor="examiner", action="export_artifact",
                        case_id=session.get("case_id"),
                        params={"artifact_id": (body.artifact or {}).get("id"), "format": ext})
    return _file_response(content, media, f"artefakt.{ext}")


@app.get("/api/session/{session_id}/export/case")
def export_session_case(session_id: str, format: str = "zip"):
    """
    Preuzmi ceo slučaj kao paket. Ako je sesija vezana za akvizicioni slučaj
    (fs_case_id), pakuje se ceo folder na disku. Inače se pravi paket sa
    izveštajima (sva 4 formata) + JSON rezultata.
    """
    session = get_session(session_id)
    fs_case_id = session.get("fs_case_id")
    if fs_case_id and acq_cases.read_case_meta(fs_case_id):
        # Osveži izveštaje u Reports/ pre pakovanja
        try:
            for fmt in ("pdf", "docx", "html", "txt"):
                content, _m, ext = exporters_mod.render(_build_export_model(session, "evidence"), fmt)
                out = acq_cases.case_dir(fs_case_id) / "Reports" / f"Analysis_Report.{ext}"
                out.parent.mkdir(parents=True, exist_ok=True)
                out.write_bytes(content if isinstance(content, bytes) else content.encode("utf-8"))
        except Exception:
            pass
        data, fname = packager_mod.build_case_zip(fs_case_id)
        audit_log.log_event(actor="examiner", action="export_case_package",
                            case_id=session.get("case_id"), params={"fs_case_id": fs_case_id})
        return _file_response(data, "application/zip", fname)

    # Nema FS slučaja (ručni dump) → paket od izveštaja + rezultata
    import io as _io, zipfile as _zip, json as _json
    buf = _io.BytesIO()
    with _zip.ZipFile(buf, "w", _zip.ZIP_DEFLATED) as zf:
        try:
            zf.writestr("Reports/Full_Report.pdf", _build_pdf_report(session))
            zf.writestr("Reports/Full_Report.docx", _build_docx_report(session))
            zf.writestr("Reports/Full_Report.html", ForensicReportEngine(_collect_report_data(session)).render())
            zf.writestr("Reports/Full_Report.txt", generate_text_report(session))
        except Exception:
            pass
        zf.writestr("Analysis/results.json", _json.dumps(session.get("results", {}), ensure_ascii=False, indent=2))
        zf.writestr("case.json", _json.dumps(_session_case_meta(session), ensure_ascii=False, indent=2))
    audit_log.log_event(actor="examiner", action="export_case_package",
                        case_id=session.get("case_id"), params={"kind": "analysis_only"})
    tag = (session.get("fs_case_id") or session.get("case_id") or "slucaj")
    return _file_response(buf.getvalue(), "application/zip", f"{tag}.zip")


# ─── Health check ─────────────────────────────────────────────────────────
@app.get("/api/health")
def health():
    return {"status": "ok", "sessions": len(SESSIONS)}


# ─── Serviranje izgrađenog frontenda (jedan proces, jedan URL) ─────────────
# Ako postoji ../build (npr. posle 'npm run build'), FastAPI servira i UI na
# istom portu (localhost:8000). Ovim cela aplikacija radi kao JEDAN proces —
# instaler ne mora da pokreće zaseban Node server. U dev-u (bez build/) ovo
# se preskače i koristi se 'npm start' + proxy.
# Kada je aplikacija spakovana (PyInstaller), 'build' se raspakuje u sys._MEIPASS;
# inače je u ../build (dev/izvorni raspored).
if getattr(sys, "frozen", False):
    _BUNDLE = Path(getattr(sys, "_MEIPASS", Path(sys.executable).parent))
    _BUILD_DIR = _BUNDLE / "build"
else:
    _BUILD_DIR = Path(__file__).resolve().parent.parent / "build"
if _BUILD_DIR.exists() and (_BUILD_DIR / "index.html").exists():
    from fastapi.staticfiles import StaticFiles
    if (_BUILD_DIR / "static").exists():
        app.mount("/static", StaticFiles(directory=str(_BUILD_DIR / "static")), name="static")

    @app.get("/{full_path:path}")
    def _serve_spa(full_path: str):
        # /api rute su registrovane iznad (imaju prednost); ovde servira SPA.
        if full_path.startswith("api/"):
            raise HTTPException(status_code=404, detail="Not found")
        candidate = _BUILD_DIR / full_path
        if full_path and candidate.exists() and candidate.is_file():
            return FileResponse(str(candidate))
        return FileResponse(str(_BUILD_DIR / "index.html"))  # SPA fallback
