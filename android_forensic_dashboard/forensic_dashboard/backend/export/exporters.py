"""
exporters.py — Univerzalni renderer izveštaja (PDF / DOCX / HTML / TXT)
────────────────────────────────────────────────────────────────────────
Ulaz je „document model" (čist dict), izlaz je fajl u traženom formatu.
Isti model → 4 formata → potpuna konzistentnost izveštaja kroz celu aplikaciju.

Document model:
  {
    "title": str,
    "subtitle": str | None,
    "meta":  [ {"label": str, "value": str}, ... ],   # zaglavlje (info o slučaju)
    "sections": [
       {"heading": str, "type": "keyvalue", "pairs": [{"label","value"}]},
       {"heading": str, "type": "table", "columns": [str], "rows": [[cell,...]]},
       {"heading": str, "type": "paragraphs", "text": [str, ...]},
       {"heading": str, "type": "list", "items": [str, ...]},
       {"heading": str, "type": "note", "text": str},
    ],
    "footer": str | None,
  }
"""

import html as _html
import io
from datetime import datetime, timezone
from pathlib import Path

# ── PDF (reportlab) ──
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, PageBreak,
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# ── DOCX ──
from docx import Document
from docx.shared import Pt, RGBColor

FOOTER_DEFAULT = ("Ovaj izveštaj sadrži faktografske nalaze (objektivni podaci). "
                  "Pravna interpretacija i krivična kvalifikacija isključivo su u "
                  "nadležnosti organa postupka.")

# ── UTF-8 font za srpska slova (isti pristup kao main.py) ──
_PDF_FONT = "Helvetica"
_PDF_FONT_BOLD = "Helvetica-Bold"
for _cand, _bold, _alias, _alias_bold in [
    (r"C:\Windows\Fonts\arial.ttf", r"C:\Windows\Fonts\arialbd.ttf", "AFDArial", "AFDArial-Bold"),
    ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
     "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", "AFDDejaVu", "AFDDejaVu-Bold"),
    ("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
     "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf", "AFDLib", "AFDLib-Bold"),
]:
    if Path(_cand).exists():
        try:
            pdfmetrics.registerFont(TTFont(_alias, _cand))
            if Path(_bold).exists():
                pdfmetrics.registerFont(TTFont(_alias_bold, _bold))
                _PDF_FONT_BOLD = _alias_bold
            _PDF_FONT = _alias
            break
        except Exception:
            pass


def _now() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")


def _s(v) -> str:
    return "" if v is None else str(v)


# ═══════════════════════════════════════════════════════════════════════════
# TXT
# ═══════════════════════════════════════════════════════════════════════════

def render_txt(model: dict) -> str:
    W = 74
    sep, thin = "═" * W, "─" * W
    L = [sep, "  " + _s(model.get("title", "")).upper(), sep]
    if model.get("subtitle"):
        L.append("  " + _s(model["subtitle"]))
    L.append(f"  Generisano: {_now()}")
    L.append("")
    for m in model.get("meta", []):
        L.append(f"  {(_s(m.get('label'))+':').ljust(24)} {_s(m.get('value'))}")
    L.append("")
    for sec in model.get("sections", []):
        L += [thin, "  " + _s(sec.get("heading", "")).upper(), thin]
        t = sec.get("type")
        if t == "keyvalue":
            for p in sec.get("pairs", []):
                L.append(f"  {(_s(p.get('label'))+':').ljust(28)} {_s(p.get('value'))}")
        elif t == "table":
            cols = sec.get("columns", [])
            if cols:
                L.append("  " + " | ".join(_s(c) for c in cols))
                L.append("  " + "-" * (W - 2))
            for row in sec.get("rows", []):
                L.append("  " + " | ".join(_s(c) for c in row))
        elif t == "paragraphs":
            for para in sec.get("text", []):
                L.append("  " + _s(para))
                L.append("")
        elif t == "list":
            for it in sec.get("items", []):
                L.append(f"  • {_s(it)}")
        elif t == "note":
            L.append("  " + _s(sec.get("text", "")))
        L.append("")
    L += [sep, "  " + _s(model.get("footer") or FOOTER_DEFAULT), sep]
    return "\n".join(L)


# ═══════════════════════════════════════════════════════════════════════════
# HTML (print-friendly, samostalan dokument)
# ═══════════════════════════════════════════════════════════════════════════

_HTML_CSS = """
:root{--fg:#111827;--muted:#6b7280;--line:#e5e7eb;--accent:#1f2937;--hi:#b91c1c;}
*{box-sizing:border-box}body{font-family:'Segoe UI',Arial,sans-serif;color:var(--fg);
margin:0;padding:32px;line-height:1.5;background:#fff;font-size:13px}
h1{font-size:22px;margin:0 0 4px}h2{font-size:15px;margin:22px 0 8px;color:var(--accent);
border-bottom:2px solid var(--line);padding-bottom:4px}h3{font-size:12px;margin:14px 0 4px}
.sub{color:var(--muted);margin-bottom:14px}.meta{border-collapse:collapse;margin:8px 0 18px}
.meta td{padding:3px 12px 3px 0;vertical-align:top}.meta td:first-child{color:var(--muted);white-space:nowrap}
table.data{border-collapse:collapse;width:100%;margin:6px 0 14px;font-size:12px}
table.data th{background:var(--accent);color:#fff;text-align:left;padding:5px 8px;font-weight:600}
table.data td{border:1px solid var(--line);padding:4px 8px;vertical-align:top}
table.data tr:nth-child(even) td{background:#f9fafb}ul{margin:4px 0 14px 18px}
.note{color:var(--muted);font-style:italic;margin:6px 0}.footer{margin-top:26px;
border-top:1px solid var(--line);padding-top:10px;color:var(--muted);font-size:11px}
@media print{body{padding:0}}
"""


def render_html(model: dict) -> str:
    e = _html.escape
    parts = ["<!DOCTYPE html><html lang='sr'><head><meta charset='utf-8'>",
             f"<title>{e(_s(model.get('title')))}</title><style>{_HTML_CSS}</style></head><body>"]
    parts.append(f"<h1>{e(_s(model.get('title')))}</h1>")
    if model.get("subtitle"):
        parts.append(f"<div class='sub'>{e(_s(model['subtitle']))}</div>")
    parts.append(f"<div class='sub'>Generisano: {e(_now())}</div>")
    if model.get("meta"):
        parts.append("<table class='meta'>")
        for m in model["meta"]:
            parts.append(f"<tr><td>{e(_s(m.get('label')))}</td><td>{e(_s(m.get('value')))}</td></tr>")
        parts.append("</table>")
    for sec in model.get("sections", []):
        parts.append(f"<h2>{e(_s(sec.get('heading')))}</h2>")
        t = sec.get("type")
        if t == "keyvalue":
            parts.append("<table class='meta'>")
            for p in sec.get("pairs", []):
                parts.append(f"<tr><td>{e(_s(p.get('label')))}</td><td>{e(_s(p.get('value')))}</td></tr>")
            parts.append("</table>")
        elif t == "table":
            parts.append("<table class='data'><thead><tr>")
            for c in sec.get("columns", []):
                parts.append(f"<th>{e(_s(c))}</th>")
            parts.append("</tr></thead><tbody>")
            for row in sec.get("rows", []):
                parts.append("<tr>" + "".join(f"<td>{e(_s(c))}</td>" for c in row) + "</tr>")
            parts.append("</tbody></table>")
        elif t == "paragraphs":
            for para in sec.get("text", []):
                parts.append(f"<p>{e(_s(para))}</p>")
        elif t == "list":
            parts.append("<ul>" + "".join(f"<li>{e(_s(it))}</li>" for it in sec.get("items", [])) + "</ul>")
        elif t == "note":
            parts.append(f"<div class='note'>{e(_s(sec.get('text')))}</div>")
    parts.append(f"<div class='footer'>{e(_s(model.get('footer') or FOOTER_DEFAULT))}</div>")
    parts.append("</body></html>")
    return "".join(parts)


# ═══════════════════════════════════════════════════════════════════════════
# PDF
# ═══════════════════════════════════════════════════════════════════════════

def render_pdf(model: dict) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm,
                            leftMargin=2 * cm, rightMargin=2 * cm,
                            title=_s(model.get("title")))
    ss = getSampleStyleSheet()
    F, FB = _PDF_FONT, _PDF_FONT_BOLD
    st_title = ParagraphStyle("T", parent=ss["Title"], fontName=FB, fontSize=18, spaceAfter=4)
    st_h2 = ParagraphStyle("H2", parent=ss["Heading2"], fontName=FB, fontSize=13,
                           spaceBefore=14, spaceAfter=6, textColor=colors.HexColor("#1f2937"))
    st_body = ParagraphStyle("B", parent=ss["Normal"], fontName=F, fontSize=9, leading=13)
    st_small = ParagraphStyle("S", parent=ss["Normal"], fontName=F, fontSize=7.5, leading=10)
    st_muted = ParagraphStyle("M", parent=st_body, fontName=F, fontSize=8,
                              textColor=colors.HexColor("#6b7280"))
    esc = _html.escape

    el = [Paragraph(esc(_s(model.get("title"))), st_title)]
    if model.get("subtitle"):
        el.append(Paragraph(esc(_s(model["subtitle"])), st_muted))
    el.append(Paragraph(f"Generisano: {esc(_now())}", st_muted))
    el.append(Spacer(1, 8))

    if model.get("meta"):
        rows = [[Paragraph(esc(_s(m.get("label"))), st_small),
                 Paragraph(esc(_s(m.get("value"))), st_small)] for m in model["meta"]]
        t = Table(rows, colWidths=[4.5 * cm, 12.5 * cm])
        t.setStyle(TableStyle([
            ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#6b7280")),
            ("LINEBELOW", (0, 0), (-1, -1), 0.4, colors.HexColor("#eef1f5")),
            ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3)]))
        el.append(t)
        el.append(Spacer(1, 6))

    for sec in model.get("sections", []):
        el.append(Paragraph(esc(_s(sec.get("heading"))), st_h2))
        t = sec.get("type")
        if t == "keyvalue":
            rows = [[Paragraph(esc(_s(p.get("label"))), st_small),
                     Paragraph(esc(_s(p.get("value"))), st_small)] for p in sec.get("pairs", [])]
            if rows:
                tbl = Table(rows, colWidths=[5.5 * cm, 11.5 * cm])
                tbl.setStyle(TableStyle([
                    ("LINEBELOW", (0, 0), (-1, -1), 0.4, colors.HexColor("#f3f4f6")),
                    ("TOPPADDING", (0, 0), (-1, -1), 2), ("BOTTOMPADDING", (0, 0), (-1, -1), 2)]))
                el.append(tbl)
        elif t == "table":
            cols = sec.get("columns", [])
            data_rows = sec.get("rows", [])
            head = [Paragraph(f"<b>{esc(_s(c))}</b>", st_small) for c in cols]
            body = [[Paragraph(esc(_s(c)), st_small) for c in row] for row in data_rows]
            n = max(1, len(cols))
            width = (17.0 / n) * cm
            tbl = Table([head] + body, colWidths=[width] * n, repeatRows=1)
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f2937")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#e5e7eb")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f9fafb")]),
                ("TOPPADDING", (0, 0), (-1, -1), 2), ("BOTTOMPADDING", (0, 0), (-1, -1), 2)]))
            el.append(tbl)
        elif t == "paragraphs":
            for para in sec.get("text", []):
                el.append(Paragraph(esc(_s(para)), st_body))
                el.append(Spacer(1, 3))
        elif t == "list":
            for it in sec.get("items", []):
                el.append(Paragraph("• " + esc(_s(it)), st_body))
        elif t == "note":
            el.append(Paragraph(esc(_s(sec.get("text"))), st_muted))

    el.append(Spacer(1, 14))
    el.append(HRFlowable(width="100%", color=colors.HexColor("#e5e7eb")))
    el.append(Paragraph(esc(_s(model.get("footer") or FOOTER_DEFAULT)), st_muted))
    doc.build(el)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════
# DOCX
# ═══════════════════════════════════════════════════════════════════════════

def render_docx(model: dict) -> bytes:
    doc = Document()
    doc.add_heading(_s(model.get("title")), level=0)
    if model.get("subtitle"):
        doc.add_paragraph(_s(model["subtitle"]))
    doc.add_paragraph(f"Generisano: {_now()}")

    if model.get("meta"):
        doc.add_heading("Informacije", level=1)
        t = doc.add_table(rows=0, cols=2)
        t.style = "Light List Accent 1"
        for m in model["meta"]:
            c = t.add_row().cells
            c[0].text = _s(m.get("label")); c[1].text = _s(m.get("value"))

    for sec in model.get("sections", []):
        doc.add_heading(_s(sec.get("heading")), level=1)
        typ = sec.get("type")
        if typ == "keyvalue":
            t = doc.add_table(rows=0, cols=2)
            t.style = "Light List Accent 1"
            for p in sec.get("pairs", []):
                c = t.add_row().cells
                c[0].text = _s(p.get("label")); c[1].text = _s(p.get("value"))
        elif typ == "table":
            cols = sec.get("columns", [])
            t = doc.add_table(rows=1, cols=max(1, len(cols)))
            t.style = "Light Grid Accent 1"
            for i, c in enumerate(cols):
                t.rows[0].cells[i].text = _s(c)
            for row in sec.get("rows", []):
                cells = t.add_row().cells
                for i, val in enumerate(row):
                    if i < len(cells):
                        cells[i].text = _s(val)
        elif typ == "paragraphs":
            for para in sec.get("text", []):
                doc.add_paragraph(_s(para))
        elif typ == "list":
            for it in sec.get("items", []):
                doc.add_paragraph(_s(it), style="List Bullet")
        elif typ == "note":
            p = doc.add_paragraph()
            r = p.add_run(_s(sec.get("text"))); r.italic = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(_s(model.get("footer") or FOOTER_DEFAULT)); r.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════
# Dispečer
# ═══════════════════════════════════════════════════════════════════════════

_MEDIA = {
    "pdf":  ("application/pdf", "pdf"),
    "docx": ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"),
    "html": ("text/html; charset=utf-8", "html"),
    "txt":  ("text/plain; charset=utf-8", "txt"),
}


def render(model: dict, fmt: str = "pdf"):
    """Vrati (sadržaj_bytes_ili_str, media_type, ekstenzija) za dati format."""
    fmt = (fmt or "pdf").lower()
    if fmt == "pdf":
        content = render_pdf(model)
    elif fmt == "docx":
        content = render_docx(model)
    elif fmt == "html":
        content = render_html(model)
    elif fmt in ("txt", "text"):
        content = render_txt(model); fmt = "txt"
    else:
        raise ValueError(f"Nepoznat format izvoza: {fmt}")
    media, ext = _MEDIA[fmt]
    return content, media, ext


# ═══════════════════════════════════════════════════════════════════════════
# Graditelji document modela iz aplikacionih prikaza (view → model)
# ═══════════════════════════════════════════════════════════════════════════

def _case_meta(case_meta: dict) -> list:
    case_meta = case_meta or {}
    out = []
    for label, key in [("Slučaj (ID)", "case_id"), ("Broj slučaja", "case_number"),
                       ("Veštak", "examiner"), ("Izvor dokaza", "source"),
                       ("Uređaj", "device"), ("Dump putanja", "dump_path"),
                       ("Otvoreno", "created_at")]:
        v = case_meta.get(key)
        if v:
            out.append({"label": label, "value": v})
    return out


def model_from_timeline(events: list, case_meta: dict = None, title="Vremenska linija") -> dict:
    rows = [[(_s(e.get("ts")).replace("T", " ").replace("Z", "")),
             _s(e.get("type")).upper(), _s(e.get("value"))[:90],
             _s(e.get("source"))] for e in (events or [])]
    return {"title": title, "subtitle": f"{len(rows)} događaja",
            "meta": _case_meta(case_meta),
            "sections": [{"heading": "Hronologija", "type": "table",
                          "columns": ["Vreme", "Tip", "Vrednost", "Izvor"], "rows": rows}]}


def model_from_correlations(correlations: list, case_meta: dict = None) -> dict:
    secs = []
    for c in (correlations or []):
        pairs = [
            {"label": "Skor / pouzdanost", "value": f"{c.get('score','?')}/100 ({c.get('band') or c.get('confidence','')})"},
            {"label": "Izvori", "value": " + ".join(c.get("sources", []))},
            {"label": "Opis", "value": c.get("detail", "")},
        ]
        secs.append({"heading": f"[{c.get('id','')}] {c.get('title','')}", "type": "keyvalue", "pairs": pairs})
    if not secs:
        secs = [{"heading": "Korelacije", "type": "note", "text": "Nema pronađenih korelacija."}]
    return {"title": "Korelacije između izvora", "subtitle": f"{len(correlations or [])} korelacija",
            "meta": _case_meta(case_meta), "sections": secs}


def model_from_artifacts(module: str, data: dict, case_meta: dict = None) -> dict:
    data = data or {}
    findings = data.get("findings") or []
    arts = data.get("artifacts") or []
    alerts = data.get("alerts") or []
    sections = []
    if findings:
        sections.append({"heading": "Nalazi", "type": "keyvalue",
                         "pairs": [{"label": f.get("key"), "value": f.get("value")} for f in findings]})
    if alerts:
        sections.append({"heading": f"Upozorenja ({len(alerts)})", "type": "list", "items": alerts})
    rows = [[_s(a.get("type")).upper(), _s(a.get("value"))[:90], _s(a.get("source")),
             _s((a.get("ts") or "")).replace("T", " ").replace("Z", "")] for a in arts]
    sections.append({"heading": f"Artefakti ({len(arts)})", "type": "table",
                     "columns": ["Tip", "Vrednost", "Izvor", "Vreme"], "rows": rows})
    return {"title": f"Modul: {module}", "subtitle": f"{len(arts)} artefakata",
            "meta": _case_meta(case_meta), "sections": sections}


def model_from_single_artifact(artifact: dict, case_meta: dict = None) -> dict:
    a = artifact or {}
    pairs = [{"label": k, "value": a.get(k)} for k in
             ("type", "value", "source", "ts", "module", "id", "confidence") if a.get(k) is not None]
    hs = a.get("hash_set") or {}
    for algo in ("md5", "sha1", "sha256"):
        if hs.get(algo):
            pairs.append({"label": algo.upper(), "value": hs[algo]})
    extra = a.get("extra") or {}
    extra_pairs = [{"label": k, "value": v} for k, v in extra.items() if v not in (None, "", [])]
    sections = [{"heading": "Artefakt", "type": "keyvalue", "pairs": pairs}]
    if extra_pairs:
        sections.append({"heading": "Detalji (extra)", "type": "keyvalue", "pairs": extra_pairs})
    prov = a.get("raw_source") or {}
    if prov:
        sections.append({"heading": "Poreklo (provenance)", "type": "keyvalue",
                         "pairs": [{"label": k, "value": v} for k, v in prov.items()]})
    return {"title": "Artefakt — detalj", "subtitle": _s(a.get("value"))[:100],
            "meta": _case_meta(case_meta), "sections": sections}


def model_from_evidence(items: list, case_meta: dict = None, title="Evidence pregled") -> dict:
    rows = [[_s(it.get("module")), _s(it.get("type")).upper(), _s(it.get("value"))[:80],
             _s(it.get("source"))] for it in (items or [])]
    return {"title": title, "subtitle": f"{len(rows)} stavki",
            "meta": _case_meta(case_meta),
            "sections": [{"heading": "Dokazi", "type": "table",
                          "columns": ["Modul", "Tip", "Vrednost", "Izvor"], "rows": rows}]}
